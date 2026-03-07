import os
import datetime
import tempfile
import platform
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

# 한글 폰트 설정: Windows=맑은 고딕, Linux=Noto Sans CJK KR (설치 필요: packages.txt)
def _set_korean_font():
    if platform.system() == "Windows":
        matplotlib.rcParams["font.family"] = "Malgun Gothic"
        return
    # Linux: Noto CJK 폰트 파일 직접 탐색
    import matplotlib.font_manager as fm
    noto_candidates = [f for f in fm.findSystemFonts() if "NotoSansCJK" in f or "NotoSerifCJK" in f or "noto" in f.lower()]
    if noto_candidates:
        prop = fm.FontProperties(fname=noto_candidates[0])
        matplotlib.rcParams["font.family"] = prop.get_name()
    else:
        # 폰트 없으면 영문 레이블 사용 (chart 제목/축은 영문으로 이미 변경됨)
        matplotlib.rcParams["font.family"] = "DejaVu Sans"

_set_korean_font()
matplotlib.rcParams["axes.unicode_minus"] = False

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import utils.market_data as md

try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False


FONT = "맑은 고딕"
CLR_H1 = RGBColor(0x1F, 0x4E, 0x79)       # Navy #1F4E79
CLR_H2 = RGBColor(0x2E, 0x75, 0xB6)       # Blue #2E75B6
CLR_LINK = "2196F3"
CLR_PAGE_NUM = "9E9E9E"
CLR_HEADER_BG = "1F4E79"
CLR_ALT_ROW = "F2F7FB"

# ================= DOCX Utilities =================

def add_hyperlink(paragraph, url: str, text: str, size_pt: float = 10):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), CLR_LINK)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    twips = str(int(size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), twips)
        rPr.append(el)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)

def add_toc(doc):
    from docx.shared import Mm as _Mm
    for lvl, indent in [("TOC 1", _Mm(0)), ("TOC 2", _Mm(5))]:
        try:
            toc_style = doc.styles[lvl]
        except KeyError:
            toc_style = doc.styles.add_style(lvl, 1)
        toc_style.font.size = Pt(10)
        toc_style.font.name = FONT
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)
        toc_style.paragraph_format.line_spacing = Pt(11)
        toc_style.paragraph_format.left_indent = indent
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    def _fld_run(fld_type=None, instr=None):
        r = p.add_run()
        if fld_type:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_type)
            r._element.append(fc)
        if instr:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            r._element.append(it)
        return r
    _fld_run("begin")
    _fld_run(instr=' TOC \\o "1-2" \\h \\z \\u ')
    _fld_run("separate")
    placeholder = p.add_run("목차 갱신: Ctrl+A → F9")
    placeholder.italic = True
    placeholder.font.size = Pt(9)
    placeholder.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    _fld_run("end")
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)

def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    def _r(fld_type=None, instr=None, literal=None):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for tag, val in [("w:sz", "18"), ("w:szCs", "18")]:
            el = OxmlElement(tag); el.set(qn("w:val"), val); rPr.append(el)
        clr = OxmlElement("w:color"); clr.set(qn("w:val"), CLR_PAGE_NUM); rPr.append(clr)
        r.append(rPr)
        if fld_type:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), fld_type); r.append(fc)
        if instr:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
            it.text = f" {instr} "; r.append(it)
        if literal:
            t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
            t.text = literal; r.append(t)
        return r
    p = para._p
    for item in [
        _r("begin"), _r(instr="PAGE"), _r("separate"), _r(literal="1"), _r("end"),
        _r(literal=" / "),
        _r("begin"), _r(instr="NUMPAGES"), _r("separate"), _r(literal="1"), _r("end"),
    ]:
        p.append(item)

def _setup_doc(title=None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = CLR_H1
    h1.font.name = FONT
    h1.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h1.paragraph_format.page_break_before = False
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = CLR_H2
    h2.font.name = FONT
    h2.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    try:
        h3 = doc.styles["Heading 3"]
    except KeyError:
        h3 = doc.styles.add_style("Heading 3", 1)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = CLR_H2
    h3.font.name = FONT
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(25)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(20)
    sec.header_distance = Mm(12)
    sec.footer_distance = Mm(12)
    if title:
        header = sec.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(title)
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        hr.font.name = FONT
    return doc

def _styled_table(doc, headers, rows, col_widths_mm=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(160 / 25.4 * 1440)))
    tblPr.append(tblW)
    if col_widths_mm is None:
        col_widths_mm = [160 / n_cols] * n_cols
    for i, w in enumerate(col_widths_mm):
        tw = int(w / 25.4 * 1440)
        for row_idx in range(1 + len(rows)):
            cell = tbl.cell(row_idx, i)
            cell.width = Emu(tw * 635)
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = FONT
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tcPr = cell._element.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._element.insert(0, tcPr)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), CLR_HEADER_BG)
        tcPr.append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(12)
            r.font.name = FONT
            if ri % 2 == 1:
                tcPr = cell._element.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._element.insert(0, tcPr)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), CLR_ALT_ROW)
                tcPr.append(shading)
    return tbl

def _add_news_section(doc, category, max_items=5):
    feed = md.fetch_rss_feed(category, max_items=max_items)
    news = feed.get("items", [])
    doc.add_heading(f"[{category}] 관련 최신 뉴스", level=2)
    if not news:
        doc.add_paragraph(
            "⚠️ 실시간 뉴스를 가져오지 못했습니다. (네트워크 접근 제한 또는 일시적 오류일 수 있습니다.)",
            style="Normal"
        )
        return []
    for item in news[:max_items]:
        p = doc.add_paragraph(style="List Bullet")
        title = item.get("title", "")
        link = item.get("link", "")
        if link:
            add_hyperlink(p, link, title)
        else:
            r = p.add_run(title)
            r.font.size = Pt(12)
            r.font.name = FONT
    return news

def _add_chart_to_doc(doc, fig, width_inches=5.8):
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp.close()
    try:
        fig.savefig(tmp.name, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        doc.add_picture(tmp.name, width=Inches(width_inches))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    finally:
        try:
            os.unlink(tmp.name)
        except Exception:
            pass

def _chart_growth():
    yr = md.LATEST_ACTUAL_YEAR
    names = [md.REGIONAL_DATA[r]["name_en"] for r in md.REGIONS]
    gwh = [md.REGIONAL_DATA[r]["installed_gwh"].get(yr, 0) for r in md.REGIONS]
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(names, gwh, color=["#1F4E79", "#2E75B6", "#5B9BD5", "#A5C8E1", "#D6E4F0", "#F0C27B", "#E8856C"])
    ax.set_ylabel("Installed Capacity (GWh)")
    ax.set_title(f"BESS Installed Capacity by Market ({yr})", fontsize=13, fontweight="bold")
    for bar, val in zip(bars, gwh):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, f"{val}", ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig

def _chart_region():
    yr = md.LATEST_ACTUAL_YEAR
    names = [md.REGIONAL_DATA[r]["name_en"] for r in md.REGIONS]
    gwh = [md.REGIONAL_DATA[r]["installed_gwh"].get(yr, 0) for r in md.REGIONS]
    fig, ax = plt.subplots(figsize=(7, 5))
    colors = ["#1F4E79", "#2E75B6", "#5B9BD5", "#A5C8E1", "#D6E4F0", "#F0C27B", "#E8856C"]
    ax.pie(gwh, labels=names, autopct="%1.1f%%", colors=colors, startangle=90)
    ax.set_title(f"BESS Market Share by Region ({yr})", fontsize=13, fontweight="bold")
    fig.tight_layout()
    return fig


# ================= Main Generator =================

def generate_word_report():
    """Generates a BESS Deep Analysis Word (.docx) report."""
    now = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=9)))  # KST
    now_str = now.strftime("%Y-%m-%d")
    
    out_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "output_reports"))
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"BESS_DeepAnalysis_{now.strftime('%Y%m%d%H%M%S')}.docx")

    doc = _setup_doc()
    sec0 = doc.sections[0]

    # Cover Page
    for _ in range(4): doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("BESS 심층 시장 분석 보고서")
    tr.font.size = Pt(28)
    tr.font.bold = True
    tr.font.color.rgb = CLR_H1
    tr.font.name = FONT

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Deep Market Intelligence Analysis")
    sr.font.size = Pt(16)
    sr.font.color.rgb = CLR_H2

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(now_str).font.size = Pt(12)

    for _ in range(3): doc.add_paragraph()
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.add_run("BESS EPC AI Agent System").font.size = Pt(12)
    doc.add_page_break()

    # TOC
    doc.add_heading("목차", level=1)
    add_toc(doc)
    doc.add_page_break()
    add_page_number(sec0)

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "본 보고서는 주요 글로벌 대상 시장의 BESS(Battery Energy Storage System) 산업 동향을 심층적으로 분석합니다. "
        "분석 시점 기준 글로벌 시장의 성장률과 가격 동향, 정책 프레임워크를 조망합니다."
    )
    doc.add_heading("핵심 지표 요약", level=2)
    _yr = md.LATEST_ACTUAL_YEAR
    summary_rows = [
        [f"글로벌 {_yr}년 설치 용량", f"{md.GLOBAL_CAPACITY_GWH.get(_yr, 'N/A')} GWh"],
        [f"LFP 셀 가격 ({_yr})", f"${md.LFP_CELL_PRICE.get(_yr, 'N/A')}/kWh"],
        [f"NMC 셀 가격 ({_yr})", f"${md.NMC_CELL_PRICE.get(_yr, 'N/A')}/kWh"],
        [f"시스템 CAPEX ({_yr})", f"${md.SYSTEM_CAPEX.get(_yr, 'N/A')}/kWh"],
        ["분석 시점", now_str],
    ]
    _styled_table(doc, ["항목", "값"], summary_rows, col_widths_mm=[70, 90])

    # 2. 시장별 심층 분석
    doc.add_page_break()
    doc.add_heading("2. 시장별 심층 분석", level=1)
    for idx, r_name in enumerate(md.REGIONS):
        r_data = md.REGIONAL_DATA[r_name]
        doc.add_heading(f"2.{idx+1} {r_name} ({r_data['name_en']})", level=2)
        
        detail_rows = [
            [f"설치 용량 ({_yr})", f"{r_data['installed_gwh'].get(_yr, 'N/A')} GWh"],
            ["성장률", f"{r_data['growth_rate_pct']}%"],
            ["파이프라인 (허가/계획)", f"{r_data['pipeline_gwh']} GWh"],
            ["매출 모델", r_data['revenue_model']],
        ]
        _styled_table(doc, ["항목", "데이터"], detail_rows, col_widths_mm=[55, 105])

        doc.add_heading("정책 환경 및 드라이버", level=3)
        for p in r_data['policy']:
            doc.add_paragraph(p, style="List Bullet")

        _add_news_section(doc, f"{r_name} 시장", 3)

    # 3. 주요 카테고리별 전문 동향
    doc.add_page_break()
    doc.add_heading("3. 전문 카테고리 분석", level=1)
    doc.add_paragraph("BESS 사업에 영향을 미치는 주요 거시적 및 미시적 카테고리 이슈 현황입니다.")
    cats = ["배터리 가격", "프로젝트", "경쟁사", "공급망", "정책·규제"]
    for c in cats:
        _add_news_section(doc, c, 5)

    # 4. 시각화
    doc.add_page_break()
    doc.add_heading("4. 시각화 분석", level=1)
    doc.add_heading("4.1 시장별 설치 용량", level=2)
    _add_chart_to_doc(doc, _chart_growth())
    
    doc.add_heading("4.2 지역별 시장 점유율", level=2)
    _add_chart_to_doc(doc, _chart_region())

    # 5. 시나리오 분석
    doc.add_page_break()
    doc.add_heading("5. 시나리오 분석", level=1)
    doc.add_paragraph("각국 시장 및 거시경제 상황에 따른 BESS 확산 중장기 시나리오 전망입니다.")
    scen_rows = []
    scen_rows.append(["연도", "보수적", "기준", "낙관적"])
    for yr in md.YEARS:
        cons = md.SCENARIOS["보수적 (Conservative)"]["capacity_gwh"].get(yr, 0)
        base = md.SCENARIOS["기본 (Base)"]["capacity_gwh"].get(yr, 0)
        opti = md.SCENARIOS["낙관적 (Optimistic)"]["capacity_gwh"].get(yr, 0)
        scen_rows.append([str(yr), f"{cons} GWh", f"{base} GWh", f"{opti} GWh"])
    _styled_table(doc, scen_rows[0], scen_rows[1:], col_widths_mm=[30, 40, 40, 40])

    doc.save(out_path)
    return out_path

def generate_pdf_report():
    """Generate PDF by converting the Word Deep Report via docx2pdf in an STA thread."""
    import threading
    import shutil
    import docx2pdf

    word_path = generate_word_report()
    pdf_path = word_path.replace('.docx', '.pdf')

    # COM/Word cannot handle paths with Korean (non-ASCII) characters.
    # Convert using an ASCII temp path, then copy to the final destination.
    tmp_dir = tempfile.gettempdir()
    tmp_docx = os.path.join(tmp_dir, "bess_report_tmp.docx")
    tmp_pdf  = os.path.join(tmp_dir, "bess_report_tmp.pdf")
    shutil.copy2(word_path, tmp_docx)

    errors = []
    def _convert():
        try:
            import pythoncom
            pythoncom.CoInitialize()
            try:
                docx2pdf.convert(tmp_docx, tmp_pdf)
            finally:
                pythoncom.CoUninitialize()
        except Exception as e:
            errors.append(e)

    t = threading.Thread(target=_convert, daemon=True)
    t.start()
    t.join(timeout=120)

    if errors:
        raise RuntimeError(f"PDF 변환 오류: {errors[0]}")

    if os.path.exists(tmp_pdf):
        shutil.copy2(tmp_pdf, pdf_path)
        for p in (tmp_docx, tmp_pdf):
            try:
                os.unlink(p)
            except Exception:
                pass
        return os.path.abspath(pdf_path)
    raise FileNotFoundError(f"PDF가 생성되지 않았습니다: {pdf_path}")
