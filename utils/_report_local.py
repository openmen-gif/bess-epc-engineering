import os
import datetime
import tempfile
import platform
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

# 한글 폰트 설정: Windows=맑은 고딕, Linux=Noto Sans CJK KR (설치 필요: packages.txt)
def _set_korean_font():
    if platform.system() == "Windows":
        matplotlib.rcParams["font.family"] = "Malgun Gothic"
        return
    # Linux: Noto CJK 폰트 파일 직접 탐색
    import matplotlib.font_manager as fm
    noto_candidates = [f for f in fm.findSystemFonts() if "NotoSansCJK" in f or "NotoSerifCJK" in f or "noto" in f.lower()]
    if noto_candidates:
        prop = fm.FontProperties(fname=noto_candidates[0])
        matplotlib.rcParams["font.family"] = prop.get_name()
    else:
        # 폰트 없으면 영문 레이블 사용 (chart 제목/축은 영문으로 이미 변경됨)
        matplotlib.rcParams["font.family"] = "DejaVu Sans"

_set_korean_font()
matplotlib.rcParams["axes.unicode_minus"] = False

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import re

import utils.market_data as md

# 표 데이터 셀 자동 정렬용 — 숫자/통화/퍼센트/N/A 표기는 우측 정렬, 그 외 텍스트는 좌측 정렬.
_NUMERIC_CELL_RE = re.compile(r"^[+\-]?\$?[\d,]+(\.\d+)?\s*%?$")

def _is_numeric_cell(val: str) -> bool:
    s = val.strip()
    if s in ("—", "N/A", ""):
        return True
    return bool(_NUMERIC_CELL_RE.match(s))

try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False


def _report_year_label(y):
    """연도에 실적/추정(E)/전망(F) 성격을 붙여 하드코딩 스냅샷을 정직하게 표기.
    값은 회계 실적이 아니라 분기 스냅샷 추정·전망치이므로, 과거 연도도 '추정'으로 표기한다."""
    cy = datetime.datetime.now().year
    if y < cy:
        return f"{y}년(추정)"
    if y == cy:
        return f"{y}년(E)"
    return f"{y}년(F)"


def _snap_basis():
    """스냅샷 출처·기준일 표기 문자열 (확정 실적이 아님을 명시)."""
    return f"스냅샷 {md.DATA_SNAPSHOT_AS_OF} 기준 · 출처 BNEF/IEA 등 · 회계 실적 아님"


FONT = "맑은 고딕"
CLR_H1 = RGBColor(0x1F, 0x4E, 0x79)       # Navy #1F4E79
CLR_H2 = RGBColor(0x2E, 0x75, 0xB6)       # Blue #2E75B6
CLR_LINK = "2196F3"
CLR_PAGE_NUM = "9E9E9E"
CLR_HEADER_BG = "1F4E79"
CLR_ALT_ROW = "F2F7FB"

# ================= DOCX Utilities =================

def add_hyperlink(paragraph, url: str, text: str, size_pt: float = 10):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), CLR_LINK)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    twips = str(int(size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), twips)
        rPr.append(el)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)

def _add_dotted_tab_stop(paragraph, position_mm):
    """Add a right-aligned tab stop with dot leader to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(int(position_mm * 56.7)))  # mm to twips
    tabs.append(tab)

def _add_bookmark(paragraph, bm_name):
    """Add a bookmark to a paragraph for internal linking."""
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(hash(bm_name) % 100000))
    bm_start.set(qn("w:name"), bm_name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(hash(bm_name) % 100000))
    paragraph._p.insert(0, bm_start)
    paragraph._p.append(bm_end)

def _make_hl_run(font_size, font_name, color, bold, text):
    """Create a w:r element for use inside a hyperlink."""
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.append(rFonts)
    twips = str(int(font_size * 2))
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), twips); rPr.append(el)
    color_hex = f"{color.red:02X}{color.green:02X}{color.blue:02X}" if hasattr(color, 'red') else str(color)
    clr = OxmlElement("w:color"); clr.set(qn("w:val"), color_hex); rPr.append(clr)
    if bold:
        b = OxmlElement("w:b"); b.set(qn("w:val"), "true"); rPr.append(b)
    # w:hyperlink 내부 run은 Word가 기본 "Hyperlink" 문자 스타일(파란색+밑줄)을
    # 암묵 적용하는 경우가 있어, 밑줄을 명시적으로 꺼서 일반 TOC 항목처럼(필드
    # 업데이트가 끝난 완성된 표 형태로) 보이게 한다 — 웹링크처럼 보인다는 지적 반영.
    u = OxmlElement("w:u"); u.set(qn("w:val"), "none"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run

def _make_pageref_field(bookmark_name, font_size, font_name, placeholder="…"):
    """Create a PAGEREF field that resolves to the actual page number of a bookmark.
    placeholder: 필드 업데이트 전(F9 누르기 전)에 표시될 값. 예상 페이지 번호를 넣어두면
    Word의 자동 필드 업데이트가 비활성화된 환경에서도 가독성 있게 표시됨.
    """
    # fldChar BEGIN
    fc_begin = OxmlElement("w:r")
    fc_begin_rPr = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), str(int(font_size * 2))); fc_begin_rPr.append(el)
    clr = OxmlElement("w:color"); clr.set(qn("w:val"), "808080"); fc_begin_rPr.append(clr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri"); rFonts.set(qn("w:eastAsia"), font_name)
    fc_begin_rPr.append(rFonts)
    fc_begin.append(fc_begin_rPr)
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    fc_begin.append(fld_begin)

    # instrText
    fc_instr = OxmlElement("w:r")
    fc_instr_rPr = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), str(int(font_size * 2))); fc_instr_rPr.append(el)
    clr2 = OxmlElement("w:color"); clr2.set(qn("w:val"), "808080"); fc_instr_rPr.append(clr2)
    fc_instr.append(fc_instr_rPr)
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f' PAGEREF {bookmark_name} \\h '
    fc_instr.append(instr)

    # fldChar SEPARATE
    fc_sep = OxmlElement("w:r")
    fc_sep_rPr = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), str(int(font_size * 2))); fc_sep_rPr.append(el)
    fc_sep.append(fc_sep_rPr)
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    fc_sep.append(fld_sep)

    # Placeholder text (shown before field update)
    fc_text = OxmlElement("w:r")
    fc_text_rPr = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), str(int(font_size * 2))); fc_text_rPr.append(el)
    clr3 = OxmlElement("w:color"); clr3.set(qn("w:val"), "808080"); fc_text_rPr.append(clr3)
    u3 = OxmlElement("w:u"); u3.set(qn("w:val"), "none"); fc_text_rPr.append(u3)
    fc_text.append(fc_text_rPr)
    t_el = OxmlElement("w:t"); t_el.text = str(placeholder)
    fc_text.append(t_el)

    # fldChar END
    fc_end = OxmlElement("w:r")
    fc_end_rPr = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag); el.set(qn("w:val"), str(int(font_size * 2))); fc_end_rPr.append(el)
    fc_end.append(fc_end_rPr)
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    fc_end.append(fld_end)

    return [fc_begin, fc_instr, fc_sep, fc_text, fc_end]


def _add_toc_hyperlink(paragraph, bookmark_name, title_text, page_text, font_size, font_name, color, bold=False):
    """Add a TOC entry as a single internal hyperlink: title + tab + dynamic PAGEREF page number.
    page_text: 예상 페이지 번호 — PAGEREF 필드의 placeholder로 사용 (필드 업데이트 전 표시값).
    """
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), bookmark_name)
    # Run 1: title text
    hl.append(_make_hl_run(font_size, font_name, color, bold, title_text))
    # Run 2: tab character
    tab_run = OxmlElement("w:r")
    tab_rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(font_size * 2))); tab_rPr.append(sz)
    tab_run.append(tab_rPr)
    tab_el = OxmlElement("w:tab")
    tab_run.append(tab_el)
    hl.append(tab_run)
    # Run 3+: PAGEREF field (dynamic page number), placeholder = 예상 페이지 번호
    for field_el in _make_pageref_field(bookmark_name, 10, font_name, placeholder=page_text):
        hl.append(field_el)
    paragraph._p.append(hl)

# 목차 항목 — 본문 섹션 heading과 반드시 동기화할 것 (bookmark명은 _add_bookmark(_hN, "_secN")과 일치)
_TOC_SECTIONS = [
    (1, "_sec1", "1. Executive Summary", 3),
    (1, "_sec2", "2. 시장별 심층 분석", 5),
    (2, "_sec2_1", "2.1 한국 (South Korea)", 6),
    (2, "_sec2_2", "2.2 일본 (Japan)", 7),
    (2, "_sec2_3", "2.3 미국 (United States)", 8),
    (2, "_sec2_4", "2.4 호주 (Australia)", 9),
    (2, "_sec2_5", "2.5 영국 (United Kingdom)", 10),
    (2, "_sec2_6", "2.6 EU (European Union)", 11),
    (2, "_sec2_7", "2.7 중동 (Middle East)", 12),
    (1, "_sec3", "3. 전문 카테고리 분석", 14),
    (1, "_sec4", "4. 시각화 분석", 16),
    (2, "_sec4_1", "4.1 시장별 설치 용량", 17),
    (2, "_sec4_2", "4.2 지역별 시장 점유율", 18),
    (1, "_sec5", "5. 글로벌 트렌드 통계 및 YoY 분석", 20),
    (2, "_sec5_1", "5.1 연도별 핵심 지표 종합", 21),
    (2, "_sec5_2", "5.2 배터리 셀 가격 및 시스템 CAPEX 추이", 22),
    (2, "_sec5_3", "5.3 글로벌 설치 용량 성장 추이", 23),
    (1, "_sec6", "6. 경쟁사 심층 분석", 25),
    (2, "_sec6_1", "6.1 주요 경쟁사 비교", 26),
    (2, "_sec6_2", "6.2 경쟁사 포지셔닝 맵", 27),
    (2, "_sec6_3", "6.3 유형별 SWOT 분석", 28),
    (1, "_sec7", "7. 프로젝트 파이프라인 현황", 30),
    (2, "_sec7_1", "7.1 주요 프로젝트 목록", 31),
    (2, "_sec7_2", "7.2 파이프라인 통계 분석", 32),
    (1, "_sec8", "8. 환율 및 원자재 시장 영향 분석", 34),
    (2, "_sec8_1", "8.1 주요 환율 현황", 35),
    (2, "_sec8_2", "8.2 원자재 가격 현황", 36),
    (2, "_sec8_3", "8.3 미국 BESS 운영용량 (EIA 라이브 추이)", 37),
    (1, "_sec9", "9. 시나리오 분석", 39),
    (2, "_sec9_1", "9.1 시나리오별 용량 전망 비교", 40),
    (2, "_sec9_2", "9.2 시나리오별 수치 비교", 41),
    (1, "_sec10", "10. BESS 사업 개발 및 투자 분석", 43),
    (2, "_sec10_1", "10.1 시장별 수익 스태킹(Revenue Stacking) 분석", 44),
    (2, "_sec10_2", "10.2 프로젝트 유형별 투자 경제성 비교", 45),
    (2, "_sec10_3", "10.3 Offtake 및 PPA 계약 구조 비교", 46),
    (1, "_sec11", "11. 전력시장 및 거래 동향", 48),
    (2, "_sec11_6", "11.6 BESS 전력 거래 전략 시사점", 49),
    (1, "_sec12", "12. BESS 운영 및 자산관리", 51),
    (2, "_sec12_1", "12.1 핵심 성능 지표(KPI)", 52),
    (2, "_sec12_2", "12.2 O&M 비용 추이", 53),
    (2, "_sec12_3", "12.3 에너지관리시스템(EMS) 플랫폼 비교", 54),
    (2, "_sec12_4", "12.4 배터리 열화 관리 및 수명 연장 전략", 55),
    (1, "_sec13", "13. 안전·화재 및 규제 기준", 57),
    (2, "_sec13_1", "13.1 글로벌 ESS 안전 규격 비교", 58),
    (2, "_sec13_2", "13.2 주요 ESS 화재 사고 사례 및 교훈", 59),
    (2, "_sec13_3", "13.3 안전 설계 체크리스트", 60),
    (1, "_sec14", "14. 배터리 기술 동향 및 차세대 기술", 62),
    (2, "_sec14_1", "14.1 배터리 기술 비교표", 63),
    (2, "_sec14_2", "14.2 기술별 상세 분석", 64),
    (2, "_sec14_3", "14.3 장기 에너지 저장(LDES) 시장 전망", 65),
    (1, "_sec15", "15. 인허가 및 사업 개발 프로세스", 67),
    (1, "_sec16", "16. 프로젝트 파이낸싱", 69),
    (2, "_sec16_1", "16.1 자금 조달 구조 비교", 70),
    (2, "_sec16_2", "16.2 Bankability 요건 (금융기관 심사 핵심)", 71),
    (2, "_sec16_3", "16.3 보험 요건", 72),
    (1, "_sec17", "17. EPC 계약 구조", 74),
    (2, "_sec17_1", "17.1 EPC 계약 유형 비교", 75),
    (2, "_sec17_2", "17.2 핵심 상업 조건", 76),
    (2, "_sec17_3", "17.3 BESS 프로젝트 비용 구조(Cost Breakdown)", 77),
    (1, "_sec18", "18. 전문가 종합 의견 및 전략적 시사점", 79),
    (2, "_sec18_1", "18.1 시장 전망 종합", 80),
    (2, "_sec18_2", "18.2 가격 전망 및 경제성 분석", 81),
    (2, "_sec18_3", "18.3 EPC 사업 전략 시사점", 82),
    (2, "_sec18_4", "18.4 핵심 리스크 요인", 83),
]

def add_toc(doc, levels: str = "1-3"):
    """정적 목차 — 섹션별 PAGEREF 하이퍼링크 항목으로 직접 구성.

    이전 방식(Word `{ TOC }` 자동생성 필드)은 Word가 문서를 열어 필드를
    갱신하기 전까지 항목 자체가 비어 있고, LibreOffice headless PDF 변환
    (HF Space의 PDF 생성 경로)은 이 자동 스캔을 수행하지 않아 목차가
    통째로 누락되는 문제가 있었다. PAGEREF는 북마크 하나만 단순 조회하면
    되므로 훨씬 안정적으로 해석되며, 설령 필드가 갱신되지 않아도 각 줄의
    제목 텍스트(및 예상 페이지 번호 placeholder)는 일반 텍스트로 항상 표시된다.
    """
    try:
        toc1_style = doc.styles["toc 1"]
    except KeyError:
        toc1_style = None
    try:
        toc2_style = doc.styles["toc 2"]
    except KeyError:
        toc2_style = None
    # 2레벨 목차: level 1(장)은 toc1 스타일·10pt·들여쓰기 없음,
    # level 2(절, N.M)는 toc2 스타일·9pt·6mm 들여쓰기로 하위 항목임을 시각적으로 구분.
    for level, bm_name, title, est_page in _TOC_SECTIONS:
        if level == 1:
            p = doc.add_paragraph(style=toc1_style)
            font_size, color, tab_pos = 10, CLR_H1, 150
        else:
            p = doc.add_paragraph(style=toc2_style)
            p.paragraph_format.left_indent = Mm(6)
            font_size, color, tab_pos = 9, CLR_H2, 144
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _add_dotted_tab_stop(p, tab_pos)
        _add_toc_hyperlink(
            p, bm_name, title, str(est_page),
            font_size=font_size, font_name=FONT, color=color, bold=(level == 1),
        )

    # 목차 끝 구분선
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(6)
    run_sep = p_sep.add_run("─" * 60)
    run_sep.font.size = Pt(8)
    run_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p_note = doc.add_paragraph()
    note_r = p_note.add_run(
        "※ 페이지 번호는 예상치이며, Word에서 열람 시 Ctrl+A → F9로 실제 값 갱신 가능합니다."
    )
    note_r.font.size = Pt(8); note_r.font.name = FONT
    note_r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    note_r.italic = True

def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    def _r(fld_type=None, instr=None, literal=None):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for tag, val in [("w:sz", "18"), ("w:szCs", "18")]:
            el = OxmlElement(tag); el.set(qn("w:val"), val); rPr.append(el)
        clr = OxmlElement("w:color"); clr.set(qn("w:val"), CLR_PAGE_NUM); rPr.append(clr)
        r.append(rPr)
        if fld_type:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), fld_type); r.append(fc)
        if instr:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
            it.text = f" {instr} "; r.append(it)
        if literal:
            t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
            t.text = literal; r.append(t)
        return r
    p = para._p
    for item in [
        _r("begin"), _r(instr="PAGE"), _r("separate"), _r(literal="1"), _r("end"),
        _r(literal=" / "),
        _r("begin"), _r(instr="NUMPAGES"), _r("separate"), _r(literal="1"), _r("end"),
    ]:
        p.append(item)

def _setup_doc(title=None):
    doc = Document()
    # Force field update on open (PAGEREF page numbers)
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = CLR_H1
    h1.font.name = FONT
    h1.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h1.paragraph_format.page_break_before = False
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = CLR_H2
    h2.font.name = FONT
    h2.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h2.paragraph_format.page_break_before = False
    try:
        h3 = doc.styles["Heading 3"]
    except KeyError:
        h3 = doc.styles.add_style("Heading 3", 1)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = CLR_H2
    h3.font.name = FONT

    # TOC 1/2/3 스타일을 컴팩트하게 미리 정의 — Word 기본값(여유 큰 줄간격) 덮어쓰기
    # Word built-in 스타일과 매핑되려면:
    #  - styleId: "TOC1" / "TOC2" / "TOC3" (공백 없이)
    #  - name (w:name val): "toc 1" / "toc 2" / "toc 3" (소문자 + 공백)
    #  - w:customStyle 속성 제거 (built-in으로 표시)
    from docx.enum.style import WD_STYLE_TYPE as _WST
    _TOC_STYLE_SPEC = [
        # (style_id, name_lower, font_size_pt, indent_mm, bold, color, line_pt)
        # line_pt = 폰트 + 2pt, space_after 1pt — 컴팩트하지만 가독성 확보
        ("TOC1", "toc 1", 10, 0,  True,  CLR_H1, 12),
        ("TOC2", "toc 2", 9,  6,  False, CLR_H2, 11),
        ("TOC3", "toc 3", 9,  12, False, None,   11),
    ]
    for _sid, _name_lower, _size, _indent_mm, _bold, _color, _line_pt in _TOC_STYLE_SPEC:
        # python-docx의 styles는 display name으로 접근 — built-in이 있으면 그것을 가져옴
        try:
            _ts = doc.styles[_name_lower]      # "toc 1" 형태로 시도
        except KeyError:
            try:
                _ts = doc.styles[f"TOC {_sid[3]}"]  # 옛 대문자 이름 시도
            except KeyError:
                _ts = doc.styles.add_style(_name_lower, _WST.PARAGRAPH)

        # XML 레벨에서 customStyle 속성 제거 + styleId/name 정규화 (built-in 매핑)
        _el = _ts.element
        _el.attrib.pop(qn("w:customStyle"), None)
        _el.set(qn("w:styleId"), _sid)
        _name_el = _el.find(qn("w:name"))
        if _name_el is None:
            _name_el = OxmlElement("w:name"); _el.insert(0, _name_el)
        _name_el.set(qn("w:val"), _name_lower)

        # 폰트/색
        _ts.font.name = FONT
        _ts.font.size = Pt(_size)
        _ts.font.bold = _bold
        if _color is not None:
            _ts.font.color.rgb = _color
        if _ts.element.rPr is not None:
            _ts.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

        # 줄간격 — exact 모드로 절대값 지정 (auto + 1.0은 Word가 자체 default 사용 가능)
        _pPr = _el.find(qn("w:pPr"))
        if _pPr is None:
            _pPr = OxmlElement("w:pPr"); _el.insert(1, _pPr)
        # 기존 spacing 노드 제거 후 새로 삽입
        for _sp_old in _pPr.findall(qn("w:spacing")):
            _pPr.remove(_sp_old)
        _sp = OxmlElement("w:spacing")
        _sp.set(qn("w:line"), str(int(_line_pt * 20)))    # twips
        _sp.set(qn("w:lineRule"), "exact")                  # 절대값 강제
        _sp.set(qn("w:before"), "0")
        _sp.set(qn("w:after"), "20")                        # 1pt = 20 twips — 항목 간 살짝
        _pPr.append(_sp)

        # 들여쓰기
        if _indent_mm:
            for _in_old in _pPr.findall(qn("w:ind")):
                _pPr.remove(_in_old)
            _ind = OxmlElement("w:ind")
            _ind.set(qn("w:left"), str(int(_indent_mm / 25.4 * 1440)))  # mm → twips
            _pPr.append(_ind)

    # 여백 옵션 A(단면 대칭, bess-output-generator.md 결정트리: "디지털 배포·화면 검토용")
    # — 본 보고서는 웹 대시보드에서 다운로드하는 디지털 산출물이며 제본 계획이 없어
    # 옵션 B(좌측 제본 여유, 좌30/우20)가 아닌 옵션 A가 맞음. 좌우 비대칭 민원(2026-07-21) 반영.
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(25)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(25)
    sec.header_distance = Mm(12)
    sec.footer_distance = Mm(12)
    if title:
        header = sec.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(title)
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        hr.font.name = FONT
    return doc

def _styled_table(doc, headers, rows, col_widths_mm=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    # Force fixed layout (critical for LibreOffice PDF conversion)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)
    total_tw = int(160 / 25.4 * 1440)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_tw))
    tblPr.append(tblW)
    if col_widths_mm is None:
        col_widths_mm = [160 / n_cols] * n_cols
    # Build w:tblGrid with explicit gridCol widths
    tblGrid = tbl._tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl._tbl.insert(tbl._tbl.index(tblPr) + 1, tblGrid)
    else:
        for child in list(tblGrid):
            tblGrid.remove(child)
    col_twips = [int(w / 25.4 * 1440) for w in col_widths_mm]
    for tw in col_twips:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(tw))
        tblGrid.append(gridCol)
    # Set explicit w:tcW on every cell
    for i, tw in enumerate(col_twips):
        for row_idx in range(1 + len(rows)):
            cell = tbl.cell(row_idx, i)
            tcPr = cell._element.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._element.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(tw))
            tcW.set(qn("w:type"), "dxa")
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(12)
        r.font.name = FONT
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        tcPr = cell._element.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._element.insert(0, tcPr)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), CLR_HEADER_BG)
        tcPr.append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.cell(ri + 1, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            _val_str = str(val)
            r = p.add_run(_val_str)
            r.font.size = Pt(12)
            r.font.name = FONT
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if _is_numeric_cell(_val_str) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ri % 2 == 1:
                tcPr = cell._element.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._element.insert(0, tcPr)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), CLR_ALT_ROW)
                tcPr.append(shading)
    return tbl

def _add_news_section(doc, category, max_items=5, analysis: str = ""):
    feed = md.fetch_rss_feed(category, max_items=max_items)
    news = feed.get("items", [])
    heading_text = (f"[{category}] 분석 및 최신 뉴스" if analysis
                    else f"[{category}] 관련 최신 뉴스")
    doc.add_heading(heading_text, level=2)
    if analysis:
        p = doc.add_paragraph(analysis)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = FONT
    if not news:
        doc.add_paragraph(
            "⚠️ 실시간 뉴스를 가져오지 못했습니다. (네트워크 접근 제한 또는 일시적 오류일 수 있습니다.)",
            style="Normal"
        )
        return []
    for item in news[:max_items]:
        p = doc.add_paragraph(style="List Bullet")
        title = item.get("title", "")
        link = item.get("link", "")
        if link:
            add_hyperlink(p, link, title)
        else:
            r = p.add_run(title)
            r.font.size = Pt(12)
            r.font.name = FONT
    return news

def _add_chart_to_doc(doc, fig, width_inches=5.8):
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    tmp.close()
    try:
        fig.savefig(tmp.name, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        doc.add_picture(tmp.name, width=Inches(width_inches))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    finally:
        try:
            os.unlink(tmp.name)
        except Exception:
            pass

def _chart_growth():
    yr = md.LATEST_ACTUAL_YEAR
    names = [md.REGIONAL_DATA[r]["name_en"] for r in md.REGIONS]
    gwh = [md.REGIONAL_DATA[r]["installed_gwh"].get(yr, 0) for r in md.REGIONS]
    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.bar(names, gwh, color=["#1F4E79", "#2E75B6", "#5B9BD5", "#A5C8E1", "#D6E4F0", "#F0C27B", "#E8856C"])
    ax.set_ylabel("Installed Capacity (GWh)")
    ax.set_title(f"BESS Installed Capacity by Market ({yr})", fontsize=13, fontweight="bold")
    for bar, val in zip(bars, gwh):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, f"{val}", ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig

def _chart_region():
    yr = md.LATEST_ACTUAL_YEAR
    names = [md.REGIONAL_DATA[r]["name_en"] for r in md.REGIONS]
    gwh = [md.REGIONAL_DATA[r]["installed_gwh"].get(yr, 0) for r in md.REGIONS]
    fig, ax = plt.subplots(figsize=(7, 5))
    colors = ["#1F4E79", "#2E75B6", "#5B9BD5", "#A5C8E1", "#D6E4F0", "#F0C27B", "#E8856C"]
    ax.pie(gwh, labels=names, autopct="%1.1f%%", colors=colors, startangle=90)
    ax.set_title(f"BESS Market Share by Region ({yr})", fontsize=13, fontweight="bold")
    fig.tight_layout()
    return fig


def _chart_price_trend():
    """Line chart: LFP, NMC cell prices and System CAPEX over years."""
    years = md.YEARS
    lfp = [md.LFP_CELL_PRICE[y] for y in years]
    nmc = [md.NMC_CELL_PRICE[y] for y in years]
    capex = [md.SYSTEM_CAPEX[y] for y in years]
    fig, ax1 = plt.subplots(figsize=(8, 4.5))
    ax1.plot(years, lfp, "o-", color="#2196F3", linewidth=2, markersize=7, label="LFP Cell ($/kWh)")
    ax1.plot(years, nmc, "s-", color="#FF5722", linewidth=2, markersize=7, label="NMC Cell ($/kWh)")
    ax2 = ax1.twinx()
    ax2.plot(years, capex, "^--", color="#4CAF50", linewidth=2, markersize=7, label="System CAPEX ($/kWh)")
    ax1.set_xlabel("Year")
    ax1.set_ylabel("Cell Price ($/kWh)")
    ax2.set_ylabel("System CAPEX ($/kWh)")
    for y, v in zip(years, lfp):
        ax1.annotate(f"${v}", (y, v), textcoords="offset points", xytext=(0, 10), ha="center", fontsize=8, color="#2196F3")
    for y, v in zip(years, nmc):
        ax1.annotate(f"${v}", (y, v), textcoords="offset points", xytext=(0, -14), ha="center", fontsize=8, color="#FF5722")
    for y, v in zip(years, capex):
        ax2.annotate(f"${v}", (y, v), textcoords="offset points", xytext=(0, 10), ha="center", fontsize=8, color="#4CAF50")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper right", fontsize=9)
    ax1.set_title("Battery Cell Price & System CAPEX Trend", fontsize=13, fontweight="bold")
    ax1.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def _chart_capacity_growth():
    """Bar + line chart: Global capacity with YoY growth rate."""
    years = md.YEARS
    caps = [md.GLOBAL_CAPACITY_GWH[y] for y in years]
    yoys = [0]
    for i in range(1, len(years)):
        prev = md.GLOBAL_CAPACITY_GWH[years[i - 1]]
        cur = md.GLOBAL_CAPACITY_GWH[years[i]]
        yoys.append(round((cur - prev) / prev * 100) if prev else 0)
    fig, ax1 = plt.subplots(figsize=(8, 4.5))
    bars = ax1.bar(years, caps, color="#1F4E79", alpha=0.8, label="Capacity (GWh)")
    for bar, val in zip(bars, caps):
        ax1.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 2,
                 f"{val}", ha="center", fontsize=9, fontweight="bold")
    ax1.set_xlabel("Year")
    ax1.set_ylabel("Installed Capacity (GWh)")
    ax2 = ax1.twinx()
    ax2.plot(years[1:], yoys[1:], "ro-", linewidth=2, markersize=7, label="YoY Growth (%)")
    for y, v in zip(years[1:], yoys[1:]):
        ax2.annotate(f"+{v}%", (y, v), textcoords="offset points", xytext=(0, 10),
                     ha="center", fontsize=9, color="red")
    ax2.set_ylabel("YoY Growth (%)")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left", fontsize=9)
    ax1.set_title("Global BESS Capacity Growth & YoY Rate", fontsize=13, fontweight="bold")
    ax1.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    return fig


def _chart_competitors():
    """Scatter chart: Competitor positioning (market share vs capacity)."""
    names = [c["name"] for c in md.COMPETITORS]
    shares = [c["market_share_pct"] for c in md.COMPETITORS]
    caps = [c["capacity_gwh"] for c in md.COMPETITORS]
    revenues = [c["revenue_b_usd"] for c in md.COMPETITORS]
    countries = [c["country"] for c in md.COMPETITORS]
    color_map = {"중국": "#E53935", "한국": "#1E88E5", "미국": "#43A047", "핀란드": "#FB8C00"}
    colors = [color_map.get(co, "#757575") for co in countries]
    sizes = [r * 30 for r in revenues]
    fig, ax = plt.subplots(figsize=(8, 5))
    for i in range(len(names)):
        ax.scatter(shares[i], caps[i], s=sizes[i], c=colors[i], alpha=0.7, edgecolors="white", linewidth=1)
        ax.annotate(names[i], (shares[i], caps[i]), textcoords="offset points",
                    xytext=(8, 5), fontsize=8)
    # Legend for countries
    for co, clr in color_map.items():
        ax.scatter([], [], c=clr, s=60, label=co)
    ax.legend(loc="lower right", fontsize=9)
    ax.set_xlabel("Market Share (%)")
    ax.set_ylabel("Production Capacity (GWh)")
    ax.set_title("BESS Competitor Positioning Map", fontsize=13, fontweight="bold")
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return fig


def _chart_scenario_comparison():
    """Multi-line chart: 3 scenarios capacity projection."""
    sb = md.SCENARIOS["기본 (Base)"]
    sc = md.SCENARIOS["보수적 (Conservative)"]
    so = md.SCENARIOS["낙관적 (Optimistic)"]
    years = sorted(sb["capacity_gwh"].keys())
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.fill_between(years,
                    [sc["capacity_gwh"][y] for y in years],
                    [so["capacity_gwh"][y] for y in years],
                    alpha=0.15, color="#2196F3", label="Range")
    ax.plot(years, [sc["capacity_gwh"][y] for y in years], "v--",
            color="#FF9800", linewidth=1.5, markersize=6, label=f"Conservative (CAGR {sc['cagr_pct']}%)")
    ax.plot(years, [sb["capacity_gwh"][y] for y in years], "o-",
            color="#1F4E79", linewidth=2.5, markersize=7, label=f"Base (CAGR {sb['cagr_pct']}%)")
    ax.plot(years, [so["capacity_gwh"][y] for y in years], "^--",
            color="#4CAF50", linewidth=1.5, markersize=6, label=f"Optimistic (CAGR {so['cagr_pct']}%)")
    for y in [years[0], years[-1]]:
        for scen, clr in [(sb, "#1F4E79"), (sc, "#FF9800"), (so, "#4CAF50")]:
            val = scen["capacity_gwh"][y]
            ax.annotate(f"{val}", (y, val), textcoords="offset points",
                        xytext=(5, 5), fontsize=8, color=clr)
    ax.set_xlabel("Year")
    ax.set_ylabel("Capacity (GWh)")
    ax.set_title(f"BESS Capacity Scenario Comparison ({years[0]}-{years[-1]})", fontsize=13, fontweight="bold")
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    return fig


def _chart_fx_trend():
    """라이브 환율 월별 추이 (USD/KRW·JPY). 보고서 생성 시점에 ECB 기준환율에서 fetch.

    이 차트는 매월 실제로 변하는 '진짜 라이브' 시계열이다(스냅샷 상수가 아님).
    Returns: (fig, error_str). 데이터 없으면 (None, error).
    """
    from utils.data_fetchers import fetch_fx_timeseries
    ts = fetch_fx_timeseries(symbols=("KRW", "JPY"), months=12)
    if ts.get("error") or not ts.get("dates"):
        return None, ts.get("error") or "데이터 없음"
    dates = ts["dates"]
    krw = ts["series"].get("KRW", [])
    jpy = ts["series"].get("JPY", [])
    x = list(range(len(dates)))
    fig, ax1 = plt.subplots(figsize=(8, 4.2))
    ax1.plot(x, krw, "o-", color="#1F4E79", linewidth=2, markersize=5, label="USD/KRW")
    ax1.set_ylabel("USD/KRW", color="#1F4E79")
    ax1.tick_params(axis="y", labelcolor="#1F4E79")
    ax2 = ax1.twinx()
    ax2.plot(x, jpy, "s--", color="#E8856C", linewidth=1.8, markersize=4, label="USD/JPY")
    ax2.set_ylabel("USD/JPY", color="#E8856C")
    ax2.tick_params(axis="y", labelcolor="#E8856C")
    ax1.set_xticks(x)
    ax1.set_xticklabels([d[:7] for d in dates], rotation=45, ha="right", fontsize=7)
    if krw:
        ax1.annotate(f"{krw[-1]:,.0f}", (x[-1], krw[-1]), textcoords="offset points",
                     xytext=(0, 9), ha="center", fontsize=8, color="#1F4E79", fontweight="bold")
    ax1.set_title(f"USD/KRW·JPY Monthly Trend — ECB (as of {ts['as_of']})", fontsize=12, fontweight="bold")
    ax1.grid(axis="y", alpha=0.3)
    l1, lab1 = ax1.get_legend_handles_labels()
    l2, lab2 = ax2.get_legend_handles_labels()
    ax1.legend(l1 + l2, lab1 + lab2, loc="upper left", fontsize=8)
    fig.tight_layout()
    return fig, None


def _chart_eia_us_monthly():
    """라이브 미국 BESS 운영용량 월별 추이 (EIA-860M). EIA_API_KEY 있을 때만 동작.

    Returns: (fig, error_str). 키 없거나 실패 시 (None, error).
    """
    from utils.data_fetchers import fetch_eia_us_battery_timeseries
    ts = fetch_eia_us_battery_timeseries(months=18)
    if ts.get("error") or not ts.get("periods"):
        return None, ts.get("error") or "데이터 없음"
    periods = ts["periods"]
    gwh = ts["gwh_est"]
    x = list(range(len(periods)))
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.plot(x, gwh, "o-", color="#2E75B6", linewidth=2, markersize=5)
    ax.fill_between(x, gwh, alpha=0.12, color="#2E75B6")
    step = max(1, len(x) // 12)
    ax.set_xticks(x[::step])
    ax.set_xticklabels([periods[i] for i in x[::step]], rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("US Operating BESS (GWh, est. @4h)")
    ax.set_title(f"US BESS Operating Capacity — EIA Monthly (as of {ts['as_of']})", fontsize=12, fontweight="bold")
    ax.grid(axis="y", alpha=0.3)
    if gwh:
        ax.annotate(f"{gwh[-1]:,.1f} GWh", (x[-1], gwh[-1]), textcoords="offset points",
                    xytext=(0, 9), ha="center", fontsize=8, color="#2E75B6", fontweight="bold")
    fig.tight_layout()
    return fig, None


# ================= Main Generator =================

def generate_word_report():
    """Generates a BESS Deep Analysis Word (.docx) report."""
    now = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=9)))  # KST
    now_str = now.strftime("%Y-%m-%d")
    
    if os.path.isdir("/data"):
        out_dir = "/data/output_reports"
    else:
        out_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "output_reports"))
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"BESS_DeepAnalysis_{now.strftime('%Y%m%d%H%M%S')}.docx")

    doc = _setup_doc()
    sec0 = doc.sections[0]

    # Cover Page
    for _ in range(4): doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("BESS 심층 시장 분석 보고서")
    tr.font.size = Pt(28)
    tr.font.bold = True
    tr.font.color.rgb = CLR_H1
    tr.font.name = FONT

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("Deep Market Intelligence Analysis")
    sr.font.size = Pt(16)
    sr.font.color.rgb = CLR_H2

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(now_str).font.size = Pt(12)

    basis_p = doc.add_paragraph()
    basis_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    basis_r = basis_p.add_run(
        f"데이터 기준 — 최신 실적 연도: {md.LATEST_ACTUAL_YEAR}년 / "
        f"전망 범위: {md.YEARS[0]}~{md.YEARS[-1]}년 / "
        f"스냅샷 시점: {md.DATA_SNAPSHOT_AS_OF}"
    )
    basis_r.font.size = Pt(12)
    basis_r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    basis_r.font.name = FONT

    fresh_p = doc.add_paragraph()
    fresh_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fresh_summary = md.get_data_freshness_summary()
    _live_n = len(_fresh_summary["live_sections"])
    _snap_n = len(_fresh_summary["all_sections"]) - _live_n
    fresh_r = fresh_p.add_run(
        f"실시간 갱신: {_live_n}개 섹션 · 스냅샷: {_snap_n}개 섹션 (분기별 큐레이션)"
    )
    fresh_r.font.size = Pt(9)
    fresh_r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fresh_r.font.name = FONT

    # 생성 시점 신선도 배너 — 클릭(생성)한 '그 날짜' 기준으로 동적 계산.
    # 라이브 항목(환율·EIA·뉴스)은 생성 시점에 실시간 수집, 시장 스냅샷은 경과일/갱신 권고를 표시.
    try:
        _snap_date = datetime.datetime.strptime(md.DATA_SNAPSHOT_AS_OF, "%Y-%m-%d").date()
        _age_days = (now.date() - _snap_date).days
    except Exception:
        _age_days = None
    _age_p = doc.add_paragraph()
    _age_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _age_days is None:
        _age_txt = f"생성 시각 {now_str} · 시장 스냅샷 기준일 확인 불가"
        _age_clr = RGBColor(0x80, 0x80, 0x80)
    elif _age_days <= 35:
        _age_txt = (f"✅ 생성 시각 {now_str} 기준 — 라이브 항목(환율·미국 EIA·뉴스) 실시간 수집, "
                    f"시장 스냅샷 {_age_days}일 전(최신)")
        _age_clr = RGBColor(0x2E, 0x7D, 0x32)
    else:
        _age_txt = (f"⚠️ 생성 시각 {now_str} 기준 — 시장 스냅샷이 {_age_days}일 경과(월간 갱신 권장). "
                    f"라이브 항목(환율·미국 EIA·뉴스)만 실시간 반영됨.")
        _age_clr = RGBColor(0xC6, 0x28, 0x28)
    _age_r = _age_p.add_run(_age_txt)
    _age_r.font.size = Pt(9)
    _age_r.font.bold = True
    _age_r.font.color.rgb = _age_clr
    _age_r.font.name = FONT

    for _ in range(3): doc.add_paragraph()
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.add_run("BESS EPC AI Agent System").font.size = Pt(12)

    # TOC — "목차" 제목은 일반 paragraph로 만들어 TOC \o "1-3" 수집에서 제외
    # (Heading 1 사용 시 TOC 안에 "목차" 항목 자체가 들어가는 문제 회피)
    _toc_h = doc.add_paragraph()
    _toc_h.paragraph_format.page_break_before = True
    _toc_h.paragraph_format.space_after = Pt(8)
    _toc_run = _toc_h.add_run("목차")
    _toc_run.font.size = Pt(18)
    _toc_run.font.bold = True
    _toc_run.font.color.rgb = CLR_H1
    _toc_run.font.name = FONT
    _toc_h.paragraph_format.page_break_before = True
    add_toc(doc)
    add_page_number(sec0)

    # 1. Executive Summary
    _h1 = doc.add_heading("1. Executive Summary", level=1)
    _h1.paragraph_format.page_break_before = True
    _add_bookmark(_h1, "_sec1")
    doc.add_paragraph(
        f"본 보고서는 {now_str} 기준 데이터로 작성되었으며, "
        f"최신 실적 연도는 {md.LATEST_ACTUAL_YEAR}년입니다. "
        "주요 글로벌 대상 시장의 BESS(Battery Energy Storage System) 산업 동향을 심층적으로 분석하며, "
        "글로벌 시장의 성장률·가격 동향·정책 프레임워크를 조망합니다."
    )
    doc.add_heading("핵심 지표 요약", level=2)
    _yr = md.LATEST_ACTUAL_YEAR
    _cy2 = datetime.datetime.now().year
    summary_rows = [
        [f"글로벌 {_yr}년 도입 (최신 실적)", f"{md.GLOBAL_CAPACITY_GWH.get(_yr, 'N/A')} GWh"],
        [f"글로벌 {_cy2}년 도입 (E·전망)", f"{md.GLOBAL_CAPACITY_GWH.get(_cy2, 'N/A')} GWh"],
        [f"LFP 셀 가격 ({_yr})", f"${md.LFP_CELL_PRICE.get(_yr, 'N/A')}/kWh"],
        [f"LFP 셀 가격 ({_cy2}, E·전망)", f"${md.LFP_CELL_PRICE.get(_cy2, 'N/A')}/kWh"],
        [f"NMC 셀 가격 ({_yr})", f"${md.NMC_CELL_PRICE.get(_yr, 'N/A')}/kWh"],
        [f"NMC 셀 가격 ({_cy2}, E·전망)", f"${md.NMC_CELL_PRICE.get(_cy2, 'N/A')}/kWh"],
        [f"시스템 CAPEX ({_yr})", f"${md.SYSTEM_CAPEX.get(_yr, 'N/A')}/kWh"],
        [f"시스템 CAPEX ({_cy2}, E·전망)", f"${md.SYSTEM_CAPEX.get(_cy2, 'N/A')}/kWh"],
        ["분석 시점", now_str],
        ["데이터 스냅샷 기준", md.DATA_SNAPSHOT_AS_OF],
    ]
    _styled_table(doc, ["항목", "값"], summary_rows, col_widths_mm=[70, 90])

    # 데이터 출처 및 신선도
    doc.add_heading("데이터 출처 및 갱신 시점", level=2)
    doc.add_paragraph(
        "본 보고서의 각 섹션별 데이터 출처와 마지막 큐레이션 시점은 아래와 같습니다. "
        "환율·원자재·뉴스는 보고서 생성 시점에 실시간 fetch되며, 그 외 시장 규모·가격·지역·경쟁사 데이터는 "
        "분기별로 갱신되는 스냅샷입니다."
    )
    _freshness_rows = []
    for s in md.get_data_freshness_summary()["all_sections"]:
        _freshness_rows.append([
            s["label"],
            s["as_of"],
            s["type"],
            s["source"],
        ])
    _styled_table(
        doc,
        ["섹션", "기준일", "유형", "출처"],
        _freshness_rows,
        col_widths_mm=[50, 20, 20, 70],
    )

    # Section 1 interpretation
    _base_yr = md.YEARS[0]
    _final_yr = md.YEARS[-1]
    _lfp_base = md.LFP_CELL_PRICE.get(_base_yr, 80)
    _lfp_cur  = md.LFP_CELL_PRICE.get(_yr, _lfp_base)
    _lfp_drop = round((1 - _lfp_cur / _lfp_base) * 100)
    _cap_cur  = md.GLOBAL_CAPACITY_GWH.get(_yr, 0)
    _cap_prev = md.GLOBAL_CAPACITY_GWH.get(_yr - 1, 0)
    _yoy_g    = round((_cap_cur - _cap_prev) / _cap_prev * 100) if _cap_prev else 0
    _mkt_val  = md.GLOBAL_MARKET_VALUE_B_USD.get(_yr, 0)
    _cap_final = md.GLOBAL_CAPACITY_GWH.get(_final_yr, 0)
    _cy       = datetime.datetime.now().year
    _cap_cy   = md.GLOBAL_CAPACITY_GWH.get(_cy, 0)
    _cy_clause = (f"보고서 작성 시점인 올해({_cy}년·E)는 약 {_cap_cy} GWh 도입이 전망됩니다(BNEF). "
                  if (_cap_cy and _cy != _yr) else "")
    _p_interp = doc.add_paragraph(
        f"글로벌 BESS 시장은 {_report_year_label(_yr)} 약 {_cap_cur} GWh로 추정되며, 전년 대비 약 {_yoy_g}% 성장한 것으로 분석됩니다 "
        f"(최신 '완료 실적' 연도 기준 — {_cy}년은 진행 중이라 연간 실적이 미확정이므로 직전 연도를 기준으로 함). "
        f"{_cy_clause}"
        f"시장 규모는 약 ${_mkt_val}B USD로 추정되고, LFP 셀 단가는 {_base_yr}년 대비 약 {_lfp_drop}% 하락한 "
        f"약 ${_lfp_cur}/kWh 수준으로 추정됩니다. "
        "가격 하락·정책 지원 확대·재생에너지 연계 수요 증가가 복합적으로 시장 성장을 견인하고 있으며, "
        f"{_report_year_label(_final_yr)} {_cap_final} GWh 도달이 전망됩니다. "
        f"※ 위 수치는 {_snap_basis()}로, 확정 실적이 아닌 추정·전망치입니다."
    )
    _p_interp.paragraph_format.space_before = Pt(4)
    for _r in _p_interp.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 최근 분기 실적 콜아웃 — 현재 시점(분기 트래커) 앵커. 연간 스냅샷보다 신선한 공개치.
    # 분기 종료 후 150일(약 1개 분기+발표 유예) 초과 시 경과 경고로 전환 — RECENT_QUARTER가
    # 다음 분기로 갱신되지 않은 채 방치되는 것을 리포트 단계에서 스스로 감지한다.
    _rq = getattr(md, "RECENT_QUARTER", None)
    if _rq:
        _rq_age_days = None
        try:
            _rq_y_str, _rq_q_str = _rq["period"].split(" Q")
            _rq_end_month = int(_rq_q_str) * 3
            _rq_end = datetime.date(int(_rq_y_str), _rq_end_month, 1)
            _rq_next_month = _rq_end.replace(day=28) + datetime.timedelta(days=4)
            _rq_end = _rq_next_month - datetime.timedelta(days=_rq_next_month.day)
            _rq_age_days = (now.date() - _rq_end).days
        except Exception:
            pass
        _rq_stale = _rq_age_days is not None and _rq_age_days > 150
        _rq_global_yoy = _rq.get("global_yoy_pct")
        _rq_us_yoy = _rq.get("us_yoy_pct")
        _rq_prefix = "⚠️ [최근 분기 실적 · 경과] " if _rq_stale else "[최근 분기 실적] "
        _rq_stale_note = (
            f" 분기 종료 후 {_rq_age_days}일 경과 — market_data.RECENT_QUARTER를 다음 분기로 갱신 권장."
            if _rq_stale else ""
        )
        _p_rq = doc.add_paragraph(
            f"{_rq_prefix}{_rq['period']} 글로벌 약 {_rq['global_gwh']} GWh 도입"
            f"{f'(전년동기 +{_rq_global_yoy}%)' if _rq_global_yoy is not None else ''}, "
            f"미국 약 {_rq['us_gwh']} GWh"
            f"{f'(전년동기 +{_rq_us_yoy}%)' if _rq_us_yoy is not None else ''}, "
            f"LFP 셀 최저 ${_rq['lfp_cell_low_usd_kwh']}/kWh 조달가. "
            f"출처: {_rq['source']}. 연간 스냅샷보다 신선한 분기 트래커로, 현재 시장이 전망 경로 상에 있음을 확인.{_rq_stale_note}"
        )
        _p_rq.paragraph_format.space_before = Pt(2)
        for _r in _p_rq.runs:
            _r.font.size = Pt(9); _r.font.name = FONT
            _r.font.bold = _rq_stale
            _r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28) if _rq_stale else RGBColor(0x2E, 0x75, 0xB6)

    # 최신 기사 기반 시장 신호 — 생성 시점에 산업 뉴스 RSS를 fetch→수치 자동 추출(LLM/키 불필요).
    # "보고서 버튼 클릭 = 그 시점 공개 기사 분석"을 배포 런타임에서 구현하는 부분.
    doc.add_heading("최신 기사 기반 시장 신호 (생성 시점 자동 추출)", level=2)
    try:
        _commentary = md.extract_market_commentary(max_items=8)
    except Exception as _ce:
        _commentary = []
    if _commentary:
        _p_cdesc = doc.add_paragraph(
            "아래는 보고서를 생성하는 바로 그 시점에 산업 뉴스 RSS(Energy-Storage.News·PV-Tech·Utility Dive 등)에서 "
            "실시간 수집해 본문에서 자동 추출한 시장 수치 신호입니다. 정밀 데이터셋이 아닌 참고 신호이므로 "
            "각 원문 링크로 검증하십시오. (스냅샷 큐레이션값과 교차 확인용)"
        )
        for _r in _p_cdesc.runs:
            _r.font.size = Pt(9); _r.font.name = FONT
        for _c in _commentary[:8]:
            _figs = ", ".join(_c.get("figures", [])[:5])
            _p = doc.add_paragraph(style="List Bullet")
            if _figs:
                _rf = _p.add_run(f"[{_figs}] ")
                _rf.bold = True; _rf.font.size = Pt(9); _rf.font.name = FONT
                _rf.font.color.rgb = CLR_H2
            _title = (_c.get("title", "") or "")[:130]
            _url = _c.get("url", "")
            if _url:
                add_hyperlink(_p, _url, _title, size_pt=9)
            else:
                _rt = _p.add_run(_title); _rt.font.size = Pt(9); _rt.font.name = FONT
    else:
        _p_cnone = doc.add_paragraph(
            "ℹ️ 생성 시점에 추출된 기사 수치가 없습니다(네트워크 제한 또는 일시적 오류). "
            "환율·미국 EIA 등 다른 라이브 항목은 정상 수집됩니다."
        )
        for _r in _p_cnone.runs:
            _r.font.size = Pt(9); _r.font.name = FONT

    # 라이브 지표 스냅샷 — 보고서 생성 시각 기준 실시간 수집값 (매 생성마다 갱신)
    doc.add_heading("실시간 지표 스냅샷 (생성 시각 기준)", level=2)
    doc.add_paragraph(
        "아래 값은 보고서 생성 시점에 외부 공개 소스에서 실시간 수집되어, 위 스냅샷 전망치와 달리 "
        "보고서를 생성할 때마다 변합니다. 스냅샷(분기 큐레이션) 대비 현재값을 대조하여 표시합니다."
    )
    try:
        _live_fx = md.fetch_exchange_rates()
    except Exception as _e:
        _live_fx = {"error": str(_e)}
    _live_rows = [["지표", "현재 라이브 값", "수집 시각 / 출처", "스냅샷 대비"]]
    _fx_krw = _live_fx.get("USD_KRW")
    _fx_ts = _live_fx.get("timestamp")
    _fx_ts_str = _fx_ts.strftime("%Y-%m-%d %H:%M") if hasattr(_fx_ts, "strftime") else "—"
    if isinstance(_fx_krw, (int, float)):
        _live_rows.append(["USD/KRW 환율", f"{_fx_krw:,.1f} 원",
                           f"{_fx_ts_str} / {_live_fx.get('source', 'er-api')}",
                           "실시간 (스냅샷 아님)"])
    else:
        _live_rows.append(["USD/KRW 환율", "수집 실패",
                           f"{_live_fx.get('error', 'N/A')}", "—"])
    # EIA 미국 라이브 용량 vs 스냅샷
    try:
        from utils.data_fetchers import get_live_us_capacity_or_fallback
        _us_snap = md.REGIONAL_DATA.get("미국", {}).get("installed_gwh", {}).get(_yr, 0)
        _us_live = get_live_us_capacity_or_fallback(_us_snap)
        if _us_live["is_live"]:
            _delta = _us_live["value_gwh"] - _us_snap
            _delta_pct = (_delta / _us_snap * 100) if _us_snap else 0
            _live_rows.append(["미국 운영용량 (EIA)", f"{_us_live['value_gwh']:,.1f} GWh",
                               f"{_us_live['as_of_str']} / EIA-860M",
                               f"스냅샷 {_us_snap} GWh 대비 {_delta_pct:+.1f}%"])
        else:
            _live_rows.append(["미국 운영용량 (EIA)", f"{_us_snap} GWh (스냅샷)",
                               f"라이브 미가용: {_us_live.get('error', 'N/A')}", "—"])
    except Exception as _e:
        _live_rows.append(["미국 운영용량 (EIA)", "—", f"오류: {_e}", "—"])
    _styled_table(doc, _live_rows[0], _live_rows[1:], col_widths_mm=[38, 36, 56, 30])

    # 2. 시장별 심층 분석
    _h2 = doc.add_heading("2. 시장별 심층 분석", level=1)
    _h2.paragraph_format.page_break_before = True
    _add_bookmark(_h2, "_sec2")
    _p_sec2note = doc.add_paragraph(
        f"⚠️ 본 장의 용량·성장률·파이프라인 수치는 {_snap_basis()}입니다. "
        "연도 표기: (추정)=스냅샷 추정, (E)=당해연도 추정, (F)=전망. "
        "환율·미국(EIA) 운영용량 등 실시간 갱신 항목은 1장 '실시간 지표 스냅샷'과 8장을 참조하십시오."
    )
    for _r in _p_sec2note.runs:
        _r.font.size = Pt(9); _r.font.name = FONT
        _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    for idx, r_name in enumerate(md.REGIONS):
        r_data = md.REGIONAL_DATA[r_name]
        _h2b = doc.add_heading(f"2.{idx+1} {r_name} ({r_data['name_en']})", level=2)
        _add_bookmark(_h2b, f"_sec2_{idx+1}")

        _cur_r  = r_data["installed_gwh"].get(_yr, 0)
        _prev_r = r_data["installed_gwh"].get(_yr - 1, 0)
        _yoy_r  = round((_cur_r - _prev_r) / _prev_r * 100) if _prev_r else 0
        _kp     = ", ".join(r_data.get("key_players", [])[:5])
        _avg_sz = r_data.get("avg_project_size_mwh", "N/A")
        _scale  = "대형 그리드 스케일" if isinstance(_avg_sz, (int, float)) and _avg_sz >= 200 else "중소형"
        _p_ov = doc.add_paragraph(
            f"{r_name} 시장은 {_report_year_label(_yr)} 약 {_cur_r} GWh로 추정되며, 전년 대비 약 {_yoy_r}% 성장한 것으로 분석됩니다. "
            f"파이프라인 {r_data['pipeline_gwh']} GWh로 향후 성장 잠재력이 크며, "
            f"평균 프로젝트 규모 {_avg_sz} MWh의 {_scale} 시장이 주도합니다. "
            f"주요 참여 기업: {_kp}. (수치: {_snap_basis()})"
        )
        _p_ov.paragraph_format.space_after = Pt(4)
        for _r in _p_ov.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

        _cy_r = datetime.datetime.now().year
        _cur_r_cy = r_data["installed_gwh"].get(_cy_r, 0)
        detail_rows = [
            [f"설치 용량 ({_yr}, 추정)", f"~{_cur_r} GWh"],
            [f"설치 용량 ({_cy_r}, E·전망)", f"~{_cur_r_cy} GWh" if _cur_r_cy else "N/A"],
            [f"YoY 성장률 ({_yr - 1}→{_yr}, 추정)", f"+{_yoy_r}%"],
            ["중기 연간 성장률", f"{r_data['growth_rate_pct']}%"],
            ["파이프라인 (허가/계획)", f"{r_data['pipeline_gwh']} GWh"],
            ["주요 수익 모델", r_data['revenue_model']],
            ["데이터 성격 · 기준일", f"분기 스냅샷 추정 · {md.DATA_SNAPSHOT_AS_OF}"],
        ]
        _styled_table(doc, ["항목", "데이터"], detail_rows, col_widths_mm=[55, 105])

        doc.add_heading("주요 성장 동인", level=3)
        for kd in r_data.get("key_drivers", []):
            doc.add_paragraph(kd, style="List Bullet")

        doc.add_heading("정책 환경 및 규제", level=3)
        for p in r_data['policy']:
            doc.add_paragraph(p, style="List Bullet")

        _add_news_section(doc, f"{r_name} 시장", 3)

    # 3. 주요 카테고리별 전문 동향
    _h3 = doc.add_heading("3. 전문 카테고리 분석", level=1)
    _h3.paragraph_format.page_break_before = True
    _add_bookmark(_h3, "_sec3")
    doc.add_paragraph("BESS 사업에 영향을 미치는 주요 거시적 및 미시적 카테고리 이슈 현황입니다.")
    _yr = md.LATEST_ACTUAL_YEAR
    # CAPEX $150/kWh 진입 연도 동적 추출
    _capex_cross_for_cat = next(
        (y for y in md.YEARS if md.SYSTEM_CAPEX.get(y, 999) <= 150), None
    )
    if _capex_cross_for_cat:
        _yrs_to_threshold = max(0, _capex_cross_for_cat - _yr)
        _capex_cat_outlook = (
            f"{_capex_cross_for_cat}년경(약 {_yrs_to_threshold}년 후) $150/kWh 이하 진입이 전망됩니다."
            if _yrs_to_threshold > 0
            else f"{_yr}년 기준 이미 $150/kWh 이하에 도달하였습니다."
        )
    else:
        _capex_cat_outlook = (
            f"{md.YEARS[-1]}년 ${md.SYSTEM_CAPEX[md.YEARS[-1]]}/kWh 수준으로, "
            "$150/kWh 도달은 데이터 범위 외에서 전망됩니다."
        )
    _CAT_ANALYSIS = {
        "배터리 가격": (
            f"LFP 셀 단가는 공급 과잉과 기술 혁신으로 지속 하락하여 {_yr}년 기준 "
            f"${md.LFP_CELL_PRICE.get(_yr, 'N/A')}/kWh 수준이며, NMC는 "
            f"${md.NMC_CELL_PRICE.get(_yr, 'N/A')}/kWh입니다. "
            "중국 제조사의 규모의 경제가 가격 하락을 주도하고 있으며, "
            f"시스템 CAPEX({_yr}: ${md.SYSTEM_CAPEX.get(_yr, 'N/A')}/kWh)는 "
            f"{_capex_cat_outlook} "
            "벤더 계층별로도 가격 격차가 뚜렷합니다 — 중국 Tier 1(CATL 글로벌 점유율 약 37%, BYD)이 "
            "가격 경쟁력을 주도하고, Tier 2(Hithium 등)는 ESS 전용 특화로 최저가 구간을 형성합니다."
        ),
        "프로젝트": (
            "글로벌 BESS 프로젝트는 대형화·장시간화 추세로 4시간 이상 장기저장 비중이 확대되고 있습니다. "
            "미국·영국·호주를 중심으로 그리드 스케일 독립형(Standalone) BESS 프로젝트가 급증하며, "
            "재생에너지 연계 하이브리드 구성이 EPC 수주의 핵심 형태로 부각됩니다. "
            "PPA 기반 장기 수익 모델과 주파수 조정(FR)·용량 시장(Capacity Market) 참여가 사업성의 핵심입니다."
        ),
        "경쟁사": (
            "글로벌 BESS 시장은 CATL, BYD 등 중국 제조사와 Tesla Megapack, Fluence, Wärtsilä 등 "
            "통합 솔루션 공급자 간 경쟁이 심화되고 있습니다. "
            "CATL은 가격 경쟁력과 글로벌 점유율(약 37%)로, BYD는 수직 통합과 Blade Battery 기반 안전성으로 "
            "차별화하며, Hithium 등 중국 Tier 2 벤더는 ESS 전용 특화로 빠르게 시장을 확대하고 있습니다. "
            "EPC 관점에서는 시스템 통합 역량과 O&M 서비스 패키지가 핵심 차별화 요소이며, "
            "현지화 전략 및 프로젝트 파이낸싱 조달 역량이 수주 경쟁력을 결정합니다."
        ),
        "공급망": (
            "리튬·코발트·망간 등 핵심 광물의 공급망 안정성이 BESS 원가 리스크의 핵심 변수입니다. "
            "미국 IRA(인플레이션 감축법) 시행으로 북미 현지 생산 수요가 급증하였고, "
            "공급망 다변화를 위한 한국·일본·유럽 제조사의 현지 투자가 가속화되고 있습니다. "
            "셀 소싱 전략(단일 vs 멀티 공급사)이 프로젝트 리스크 관리의 핵심 요소로 부상하고 있습니다."
        ),
        "정책·규제": (
            "미국 IRA, 유럽 Net-Zero Industry Act, 영국 Capacity Market 등 주요 시장의 정책 지원이 "
            "BESS 수요를 견인하고 있습니다. "
            "계통 연계 기준(Grid Code) 및 안전 규정(NFPA 855, IEC 62933) 강화가 진행 중이며, "
            "ESS 화재 안전 기준과 소방 설계가 프로젝트 인허가의 핵심 요건으로 부상하고 있습니다. "
            "각국의 탄소중립 로드맵 이행이 중장기 BESS 수요 성장의 근본 동인입니다."
        ),
    }
    for c in ["배터리 가격", "프로젝트", "경쟁사", "공급망", "정책·규제"]:
        _add_news_section(doc, c, 5, analysis=_CAT_ANALYSIS.get(c, ""))

    _p_skillref = doc.add_paragraph(
        "※ 위 카테고리 해설은 실시간 뉴스 신호와 함께, 플랫폼 내 전문 스킬 문서 "
        "(skill_md/bess-battery-expert.md, skill_md/bess-marketer.md)의 벤더 계층·가격 프레임워크를 "
        "정적으로 반영합니다. 스킬 문서 갱신 시 본 섹션도 함께 재검토가 필요합니다."
    )
    for _r in _p_skillref.runs:
        _r.font.size = Pt(8); _r.font.name = FONT
        _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _r.italic = True

    # 3.X Recent Market Commentary (RSS에서 자동 추출된 시장 인용 수치)
    doc.add_heading("📊 최근 시장 인용 수치 (RSS 자동 추출)", level=2)
    doc.add_paragraph(
        "보고서 생성 시점에 fetch된 RSS 피드에서 BESS 시장 관련 수치 인용 "
        "(예: $/kWh 가격, GWh 용량, MW 출력, % 성장률) 을 자동으로 추출한 결과입니다. "
        "정량 데이터의 최신 시장 발언을 보완 자료로 활용하십시오."
    )
    try:
        _commentary = md.extract_market_commentary(max_items=10)
    except Exception as _ce:
        _commentary = []
        doc.add_paragraph(f"⚠️ 인용 추출 중 오류: {_ce}")
    if not _commentary:
        doc.add_paragraph("최근 fetch한 RSS에서 추출 가능한 수치 인용이 없습니다.")
    else:
        for _c in _commentary:
            _fig_str = ", ".join(_c["figures"]) if _c["figures"] else ""
            _p_h = doc.add_paragraph(style="List Bullet")
            _h_run = _p_h.add_run(f"[{_fig_str}] ")
            _h_run.font.bold = True
            _h_run.font.size = Pt(12)
            _h_run.font.name = FONT
            if _c.get("url"):
                add_hyperlink(_p_h, _c["url"], _c["title"])
            else:
                _t_run = _p_h.add_run(_c["title"])
                _t_run.font.size = Pt(12)
                _t_run.font.name = FONT
            _p_meta = doc.add_paragraph()
            _m_run = _p_meta.add_run(
                f"  └ {_c.get('source', '')} · {_c.get('pub_date', '')}"
            )
            _m_run.font.size = Pt(9)
            _m_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            _m_run.font.name = FONT

    # 4. 시각화
    _h4 = doc.add_heading("4. 시각화 분석", level=1)
    _h4.paragraph_format.page_break_before = True
    _add_bookmark(_h4, "_sec4")
    _yr4 = md.LATEST_ACTUAL_YEAR
    _bars = [(r, md.REGIONAL_DATA[r]["name_en"],
              md.REGIONAL_DATA[r]["installed_gwh"].get(_yr4, 0),
              md.REGIONAL_DATA[r]["growth_rate_pct"]) for r in md.REGIONS]
    _total4 = sum(x[2] for x in _bars)
    _bars_sorted = sorted(_bars, key=lambda x: x[2], reverse=True)
    _top3_sum4   = sum(x[2] for x in _bars_sorted[:3])
    _top3_pct4   = round(_top3_sum4 / _total4 * 100) if _total4 else 0
    _fastest4    = sorted(_bars, key=lambda x: x[3], reverse=True)[0]

    _h2b = doc.add_heading("4.1 시장별 설치 용량", level=2)
    _add_bookmark(_h2b, "_sec4_1")
    _add_chart_to_doc(doc, _chart_growth())
    _p_bar = doc.add_paragraph(
        f"[그림 분석] {_yr4}년 글로벌 설치 용량 합계 {_total4:.1f} GWh 중, "
        f"{_bars_sorted[0][1]}({_bars_sorted[0][2]} GWh)이 1위를 차지합니다. "
        f"상위 3개 시장({', '.join(x[1] for x in _bars_sorted[:3])})의 합산 비중은 "
        f"{_top3_pct4}%로 시장 집중도가 높습니다. "
        f"성장률 최고 시장은 {_fastest4[0]}({_fastest4[1]}, 중기 CAGR {_fastest4[3]}%)로 "
        f"신흥 시장 중 가장 빠른 확대세를 보입니다."
    )
    _p_bar.paragraph_format.space_before = Pt(4)
    for _r in _p_bar.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    _pie_sorted = sorted(_bars, key=lambda x: x[2], reverse=True)
    _shares4    = [(x[1], round(x[2] / _total4 * 100, 1)) for x in _pie_sorted]
    _top3_sh    = round(sum(s[1] for s in _shares4[:3]), 1)

    _h2b = doc.add_heading("4.2 지역별 시장 점유율", level=2)
    _add_bookmark(_h2b, "_sec4_2")
    _add_chart_to_doc(doc, _chart_region())
    _p_pie = doc.add_paragraph(
        f"[그림 분석] {_shares4[0][0]}이 {_shares4[0][1]}%로 압도적 1위이며, "
        f"{_shares4[1][0]} {_shares4[1][1]}%, {_shares4[2][0]} {_shares4[2][1]}%가 뒤를 잇습니다. "
        f"상위 3개 지역 합산 점유율 {_top3_sh}%로 시장이 집중되어 있으며, "
        f"{_shares4[-1][0]}({_shares4[-1][1]}%)는 아직 소규모이나 "
        f"중기 성장률 {_fastest4[3]}%로 가장 빠른 성장이 전망됩니다."
    )
    _p_pie.paragraph_format.space_before = Pt(4)
    for _r in _p_pie.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 5. 글로벌 트렌드 통계 및 YoY 분석
    _h5 = doc.add_heading("5. 글로벌 트렌드 통계 및 YoY 분석", level=1)
    _h5.paragraph_format.page_break_before = True
    _add_bookmark(_h5, "_sec5")
    doc.add_paragraph(
        "BESS 시장의 핵심 지표에 대한 연도별 추이 및 성장률을 정량적으로 분석합니다. "
        "용량, 시장 가치, 셀 가격, 시스템 CAPEX의 연도별 변화율(YoY)과 "
        "복합 연간 성장률(CAGR)을 산출하여 시장의 구조적 트렌드를 파악합니다."
    )

    # 5.1 연도별 핵심 지표 종합 테이블
    _h2b = doc.add_heading("5.1 연도별 핵심 지표 종합", level=2)
    _add_bookmark(_h2b, "_sec5_1")
    trend_headers = ["연도", "용량(GWh)", "YoY(%)", "시장규모($B)", "YoY(%)",
                     "LFP($/kWh)", "변동(%)", "CAPEX($/kWh)", "변동(%)"]
    trend_rows = []
    for i, yr in enumerate(md.YEARS):
        cap = md.GLOBAL_CAPACITY_GWH.get(yr, 0)
        cap_p = md.GLOBAL_CAPACITY_GWH.get(md.YEARS[i - 1], 0) if i > 0 else 0
        cap_yoy = f"+{round((cap - cap_p) / cap_p * 100)}%" if cap_p else "—"
        mkt = md.GLOBAL_MARKET_VALUE_B_USD.get(yr, 0)
        mkt_p = md.GLOBAL_MARKET_VALUE_B_USD.get(md.YEARS[i - 1], 0) if i > 0 else 0
        mkt_yoy = f"+{round((mkt - mkt_p) / mkt_p * 100)}%" if mkt_p else "—"
        lfp = md.LFP_CELL_PRICE.get(yr, 0)
        lfp_p = md.LFP_CELL_PRICE.get(md.YEARS[i - 1], 0) if i > 0 else 0
        lfp_chg = f"{round((lfp - lfp_p) / lfp_p * 100)}%" if lfp_p else "—"
        capex = md.SYSTEM_CAPEX.get(yr, 0)
        capex_p = md.SYSTEM_CAPEX.get(md.YEARS[i - 1], 0) if i > 0 else 0
        capex_chg = f"{round((capex - capex_p) / capex_p * 100)}%" if capex_p else "—"
        trend_rows.append([str(yr), str(cap), cap_yoy, f"${mkt}", mkt_yoy,
                           f"${lfp}", lfp_chg, f"${capex}", capex_chg])
    _styled_table(doc, trend_headers, trend_rows,
                  col_widths_mm=[18, 18, 14, 18, 14, 18, 14, 32, 14])

    # CAGR calculations
    _first_yr, _last_yr = md.YEARS[0], md.YEARS[-1]
    _n_yrs = _last_yr - _first_yr
    _cap_cagr = round(((md.GLOBAL_CAPACITY_GWH[_last_yr] / md.GLOBAL_CAPACITY_GWH[_first_yr])
                        ** (1 / _n_yrs) - 1) * 100, 1) if md.GLOBAL_CAPACITY_GWH[_first_yr] else 0
    _mkt_cagr = round(((md.GLOBAL_MARKET_VALUE_B_USD[_last_yr] / md.GLOBAL_MARKET_VALUE_B_USD[_first_yr])
                        ** (1 / _n_yrs) - 1) * 100, 1) if md.GLOBAL_MARKET_VALUE_B_USD[_first_yr] else 0
    _lfp_cagr = round(((md.LFP_CELL_PRICE[_last_yr] / md.LFP_CELL_PRICE[_first_yr])
                        ** (1 / _n_yrs) - 1) * 100, 1) if md.LFP_CELL_PRICE[_first_yr] else 0

    _p_cagr = doc.add_paragraph(
        f"[통계 분석] {_first_yr}~{_last_yr}년 기간 CAGR: "
        f"설치 용량 {_cap_cagr}%, 시장 규모 {_mkt_cagr}%, "
        f"LFP 셀 가격 {_lfp_cagr}% (가격 하락은 음수). "
        f"용량 기준 연평균 {_cap_cagr}% 성장은 글로벌 에너지 전환의 핵심 인프라로서 "
        f"BESS의 구조적 성장을 입증합니다. "
        f"LFP 셀 가격이 매년 약 {abs(_lfp_cagr)}%씩 하락하면서 경제성이 개선되고 있으며, "
        f"시스템 CAPEX도 {md.SYSTEM_CAPEX[_first_yr]}→{md.SYSTEM_CAPEX[_last_yr]} $/kWh로 "
        f"{round((1 - md.SYSTEM_CAPEX[_last_yr] / md.SYSTEM_CAPEX[_first_yr]) * 100)}% 감소하였습니다."
    )
    _p_cagr.paragraph_format.space_before = Pt(4)
    for _r in _p_cagr.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 5.2 가격 트렌드 차트
    _h2b = doc.add_heading("5.2 배터리 셀 가격 및 시스템 CAPEX 추이", level=2)
    _add_bookmark(_h2b, "_sec5_2")
    _add_chart_to_doc(doc, _chart_price_trend())
    _capex_threshold = 150
    _capex_cross_yr = next(
        (y for y in md.YEARS if md.SYSTEM_CAPEX.get(y, 999) <= _capex_threshold), None
    )
    _capex_outlook = (
        f"{_capex_cross_yr}년경 시스템 CAPEX ${_capex_threshold}/kWh 이하 진입이 전망됩니다."
        if _capex_cross_yr
        else f"시스템 CAPEX는 {_last_yr}년 ${md.SYSTEM_CAPEX[_last_yr]}/kWh 수준으로, "
             f"향후 ${_capex_threshold}/kWh 이하 진입은 데이터 범위 외에서 전망됩니다."
    )
    _p_price = doc.add_paragraph(
        f"[그림 분석] LFP 셀 가격은 {md.LFP_CELL_PRICE[_first_yr]}→{md.LFP_CELL_PRICE[_last_yr]} $/kWh로 "
        f"{round((1 - md.LFP_CELL_PRICE[_last_yr] / md.LFP_CELL_PRICE[_first_yr]) * 100)}% 하락하였으며, "
        f"NMC 셀도 {md.NMC_CELL_PRICE[_first_yr]}→{md.NMC_CELL_PRICE[_last_yr]} $/kWh로 동반 하락했습니다. "
        f"시스템 CAPEX(BOS 포함)는 {md.SYSTEM_CAPEX[_first_yr]}→{md.SYSTEM_CAPEX[_last_yr]} $/kWh로 "
        "LFP 대비 완만한 하락세를 보이는데, 이는 인버터·EMS·토목 비용이 셀 가격보다 "
        f"경직적이기 때문입니다. {_capex_outlook}"
    )
    _p_price.paragraph_format.space_before = Pt(4)
    for _r in _p_price.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 5.3 용량 성장 추이 차트
    _h2b = doc.add_heading("5.3 글로벌 설치 용량 성장 추이", level=2)
    _add_bookmark(_h2b, "_sec5_3")
    _add_chart_to_doc(doc, _chart_capacity_growth())
    _p_cap_g = doc.add_paragraph(
        f"[그림 분석] 글로벌 BESS 설치 용량은 {_first_yr}년 {md.GLOBAL_CAPACITY_GWH[_first_yr]} GWh에서 "
        f"{_last_yr}년 {md.GLOBAL_CAPACITY_GWH[_last_yr]} GWh로 약 "
        f"{round(md.GLOBAL_CAPACITY_GWH[_last_yr] / md.GLOBAL_CAPACITY_GWH[_first_yr], 1)}배 성장하였습니다. "
        f"CAGR {_cap_cagr}%의 초고속 성장은 재생에너지 확대, 전력 계통 유연성 수요, "
        f"배터리 가격 하락이 복합적으로 작용한 결과입니다."
    )
    _p_cap_g.paragraph_format.space_before = Pt(4)
    for _r in _p_cap_g.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 6. 경쟁사 심층 분석
    _h6 = doc.add_heading("6. 경쟁사 심층 분석", level=1)
    _h6.paragraph_format.page_break_before = True
    _add_bookmark(_h6, "_sec6")
    doc.add_paragraph(
        "글로벌 BESS 시장 주요 플레이어의 시장 점유율, 매출, 생산 역량 및 강약점을 비교 분석합니다. "
        "셀 제조사와 시스템 통합업체(SI)의 포지셔닝 차이를 이해하는 것이 EPC 수주 전략 수립의 핵심입니다."
    )

    # 6.1 시장 점유율 테이블
    _h2b = doc.add_heading("6.1 주요 경쟁사 비교", level=2)
    _add_bookmark(_h2b, "_sec6_1")
    comp_headers = ["기업", "국가", "유형", "점유율(%)", "매출($B)", "생산(GWh)"]
    comp_rows = []
    for c in md.COMPETITORS:
        comp_rows.append([
            c["name"], c["country"], c["type"],
            f"{c['market_share_pct']}%", f"${c['revenue_b_usd']}",
            str(c["capacity_gwh"])
        ])
    _styled_table(doc, comp_headers, comp_rows, col_widths_mm=[46, 16, 22, 20, 28, 28])

    # 6.2 경쟁사 포지셔닝 차트
    _h2b = doc.add_heading("6.2 경쟁사 포지셔닝 맵", level=2)
    _add_bookmark(_h2b, "_sec6_2")
    _add_chart_to_doc(doc, _chart_competitors())

    # Top 3 competitor analysis
    _sorted_comp = sorted(md.COMPETITORS, key=lambda x: x["market_share_pct"], reverse=True)
    _top3_comp = _sorted_comp[:3]
    _top3_share = sum(c["market_share_pct"] for c in _top3_comp)
    _cn_share = sum(c["market_share_pct"] for c in md.COMPETITORS if c["country"] == "중국")
    _p_comp = doc.add_paragraph(
        f"[분석] 상위 3개 기업({', '.join(c['name'] for c in _top3_comp)})의 합산 점유율은 "
        f"{_top3_share}%로, 시장 집중도가 매우 높습니다. "
        f"중국 기업의 합산 점유율은 {_cn_share}%로 절대적이며, "
        "LFP 원가 경쟁력과 대규모 생산 역량이 핵심 우위입니다. "
        "한편 Tesla Megapack, Fluence 등 미국 SI는 소프트웨어 통합 역량과 브랜드 신뢰로 "
        "서방 시장에서 강한 입지를 유지하고 있습니다."
    )
    _p_comp.paragraph_format.space_before = Pt(4)
    for _r in _p_comp.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 6.3 SWOT by player type
    _h2b = doc.add_heading("6.3 유형별 SWOT 분석", level=2)
    _add_bookmark(_h2b, "_sec6_3")
    for c in _sorted_comp[:5]:
        _p_sw = doc.add_paragraph()
        _r_name = _p_sw.add_run(f"{c['name']} ({c['country']}, {c['type']}): ")
        _r_name.bold = True
        _r_name.font.size = Pt(12); _r_name.font.name = FONT
        _r_body = _p_sw.add_run(f"강점 — {c['strength']}  |  약점 — {c['weakness']}")
        _r_body.font.size = Pt(12); _r_body.font.name = FONT

    # 7. 프로젝트 파이프라인 현황
    _h7 = doc.add_heading("7. 프로젝트 파이프라인 현황", level=1)
    _h7.paragraph_format.page_break_before = True
    _add_bookmark(_h7, "_sec7")
    doc.add_paragraph(
        "글로벌 주요 BESS 프로젝트의 규모, 상태, 지역 분포를 분석합니다. "
        "대형 프로젝트 트렌드와 주요 개발사의 활동을 통해 시장의 방향성을 파악합니다."
    )

    # 7.1 파이프라인 테이블
    _h2b = doc.add_heading("7.1 주요 프로젝트 목록", level=2)
    _add_bookmark(_h2b, "_sec7_1")
    _pipeline = md.project_pipeline_with_status()
    pipe_headers = ["프로젝트명", "지역", "MW", "MWh", "상태", "개발사", "연도"]
    pipe_rows = []
    for p in _pipeline:
        pipe_rows.append([
            p["name"], p["region"], str(p["capacity_mw"]),
            str(p["capacity_mwh"]), p["status"], p["developer"], str(p["year"])
        ])
    _styled_table(doc, pipe_headers, pipe_rows,
                  col_widths_mm=[50, 14, 14, 14, 14, 36, 18])

    # 7.2 파이프라인 통계
    _h2b = doc.add_heading("7.2 파이프라인 통계 분석", level=2)
    _add_bookmark(_h2b, "_sec7_2")
    _total_mw = sum(p["capacity_mw"] for p in _pipeline)
    _total_mwh = sum(p["capacity_mwh"] for p in _pipeline)
    _avg_mwh = round(_total_mwh / len(_pipeline))
    _by_status = {}
    _by_region = {}
    for p in _pipeline:
        _by_status[p["status"]] = _by_status.get(p["status"], 0) + 1
        _by_region[p["region"]] = _by_region.get(p["region"], 0) + p["capacity_mwh"]
    _top_region = max(_by_region, key=_by_region.get)

    pipe_stat_rows = [
        ["총 프로젝트 수", f"{len(_pipeline)}개"],
        ["총 용량 (MW / MWh)", f"{_total_mw:,} MW / {_total_mwh:,} MWh"],
        ["평균 프로젝트 규모", f"{_avg_mwh:,} MWh"],
        ["상태별 분포", ", ".join(f"{k}: {v}건" for k, v in _by_status.items())],
        ["최대 용량 지역", f"{_top_region} ({_by_region[_top_region]:,} MWh)"],
    ]
    _styled_table(doc, ["항목", "값"], pipe_stat_rows, col_widths_mm=[60, 100])

    _p_pipe = doc.add_paragraph(
        f"[분석] 글로벌 주요 파이프라인의 평균 프로젝트 규모는 {_avg_mwh:,} MWh로 "
        "대형화 추세가 뚜렷합니다. 4시간 이상 장기저장 프로젝트 비중이 증가하고 있으며, "
        f"{_top_region} 시장이 총 {_by_region[_top_region]:,} MWh로 최대 파이프라인을 보유하고 있습니다. "
        "건설중·계획중 프로젝트가 전체의 상당 부분을 차지하여 향후 2~3년 내 "
        "대규모 용량 증설이 예상됩니다."
    )
    _p_pipe.paragraph_format.space_before = Pt(4)
    for _r in _p_pipe.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 8. 환율 및 원자재 시장 영향 분석
    _h8 = doc.add_heading("8. 환율 및 원자재 시장 영향 분석", level=1)
    _h8.paragraph_format.page_break_before = True
    _add_bookmark(_h8, "_sec8")
    doc.add_paragraph(
        "BESS 프로젝트의 수익성과 원가에 직접 영향을 미치는 환율 및 원자재 가격 동향을 분석합니다."
    )

    # Try to fetch live data
    # 최후 폴백값 — market_data.fetch_commodity_prices()의 자체 참조값과 반드시 동기화할 것
    # (2026-07 재검증 기준. 갱신 시 market_data.py의 _FALLBACK/_COMMODITY_REF_DATE도 함께 갱신)
    _FX_FALLBACK_DATE = "2026-07"
    try:
        _fx = md.fetch_exchange_rates()
        _cmd = md.fetch_commodity_prices()
    except Exception:
        _fx = {"USD_KRW": 1350, "USD_JPY": 150, "USD_EUR": 0.92, "USD_CNY": 7.25,
               "USD_AUD": 1.55, "USD_GBP": 0.79,
               "error": "offline",
               "source": f"reference (as of {_FX_FALLBACK_DATE}, offline)"}
        _cmd = {"brent_crude_usd": 72.5, "wti_crude_usd": 68.8,
                "lithium_carbonate_usd_ton": 23000, "copper_usd_ton": 9200,
                "nickel_usd_ton": 16800,
                "source": f"reference (as of {_FX_FALLBACK_DATE}, offline)"}

    _h2b = doc.add_heading("8.1 주요 환율 현황", level=2)
    _add_bookmark(_h2b, "_sec8_1")
    fx_rows = [
        ["USD/KRW", f"{_fx.get('USD_KRW', 'N/A'):,.1f}" if isinstance(_fx.get('USD_KRW'), (int, float)) else "N/A",
         "한국 시장 원가에 직접 영향"],
        ["USD/JPY", f"{_fx.get('USD_JPY', 'N/A'):,.1f}" if isinstance(_fx.get('USD_JPY'), (int, float)) else "N/A",
         "일본 시장 수출 경쟁력"],
        ["USD/EUR", f"{_fx.get('USD_EUR', 'N/A'):.4f}" if isinstance(_fx.get('USD_EUR'), (int, float)) else "N/A",
         "EU 시장 입찰 환산"],
        ["USD/CNY", f"{_fx.get('USD_CNY', 'N/A'):.4f}" if isinstance(_fx.get('USD_CNY'), (int, float)) else "N/A",
         "중국 셀 수입가 결정"],
        ["USD/AUD", f"{_fx.get('USD_AUD', 'N/A'):.4f}" if isinstance(_fx.get('USD_AUD'), (int, float)) else "N/A",
         "호주 시장 수주 환산"],
    ]
    _styled_table(doc, ["통화쌍", "환율", "BESS 영향"], fx_rows, col_widths_mm=[30, 30, 100])

    # 라이브 환율 추이 차트 (매월 변하는 실시간 시계열)
    _fig_fx, _err_fx = _chart_fx_trend()
    if _fig_fx is not None:
        _add_chart_to_doc(doc, _fig_fx)
        _p_fxnote = doc.add_paragraph(
            "※ 위 환율 추이는 보고서 생성 시점에 ECB(유럽중앙은행) 기준환율에서 실시간 수집한 "
            "월별 시계열입니다. 스냅샷 상수가 아니므로 매월 보고서를 생성할 때마다 갱신됩니다."
        )
        for _r in _p_fxnote.runs:
            _r.font.size = Pt(8); _r.font.name = FONT
            _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        _p_fxerr = doc.add_paragraph(
            f"⚠️ 실시간 환율 추이 차트를 가져오지 못했습니다 ({_err_fx}). "
            "네트워크 제한 또는 일시적 오류일 수 있으며, 위 표의 현재값을 참고하십시오."
        )
        for _r in _p_fxerr.runs:
            _r.font.size = Pt(9); _r.font.name = FONT

    _h2b = doc.add_heading("8.2 원자재 가격 현황", level=2)
    _add_bookmark(_h2b, "_sec8_2")
    cmd_rows = [
        ["브렌트유 (Brent)",
         f"${_cmd.get('brent_crude_usd', 'N/A')}/bbl",
         "운송비·디젤 발전 대체 경제성에 영향"],
        ["리튬 카보네이트",
         f"${_cmd.get('lithium_carbonate_usd_ton', 'N/A'):,}/톤",
         "LFP/NMC 셀 원가의 핵심 변수 (셀 원가의 30~40%)"],
        ["구리",
         f"${_cmd.get('copper_usd_ton', 'N/A'):,}/톤",
         "케이블·버스바·인버터 원가에 직접 영향"],
        ["니켈",
         f"${_cmd.get('nickel_usd_ton', 'N/A'):,}/톤",
         "NMC 셀 원가 변수 (LFP 전환 가속 요인)"],
    ]
    _styled_table(doc, ["원자재", "가격", "BESS 산업 영향"], cmd_rows, col_widths_mm=[30, 30, 100])

    _p_fx = doc.add_paragraph(
        "[영향 분석] 원/달러 환율 상승(원화 약세)은 한국 EPC 기업의 해외 수주 가격 경쟁력을 높이나, "
        "셀·원자재 수입 비용을 증가시키는 양면적 효과가 있습니다. "
        "리튬 카보네이트 가격은 역대 최고점(2022년 11월 약 $80,000/톤) 대비 대폭 하락하여 현재 안정권이며, "
        "이는 LFP 셀 가격 하락의 핵심 동인입니다. "
        "다만 리튬 가격의 바닥 도달 시 하락 속도가 둔화될 수 있어, "
        "중기적으로 셀 가격 하락 폭이 축소될 가능성에 유의해야 합니다."
    )
    _p_fx.paragraph_format.space_before = Pt(4)
    for _r in _p_fx.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    _cmd_src = _cmd.get("metals_source") or _cmd.get("source") or "reference"
    _p_cmdsrc = doc.add_paragraph(
        f"※ 환율은 실시간(open.er-api.com) 수집값입니다. 원자재 중 리튬·구리·니켈은 무료 실시간 "
        f"공개 시세가 없어 분기 참조값({_cmd_src})이며, 브렌트유는 공개 API 가용 시 갱신됩니다. "
        "따라서 금속 가격은 월별로 변하지 않을 수 있습니다(데이터 신선도 표 참조)."
    )
    for _r in _p_cmdsrc.runs:
        _r.font.size = Pt(8); _r.font.name = FONT
        _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 8.3 미국 BESS 운영용량 — EIA 라이브 월별 추이
    _h2b = doc.add_heading("8.3 미국 BESS 운영용량 (EIA 라이브 추이)", level=2)
    _add_bookmark(_h2b, "_sec8_3")
    _fig_eia, _err_eia = _chart_eia_us_monthly()
    if _fig_eia is not None:
        _add_chart_to_doc(doc, _fig_eia)
        _p_eianote = doc.add_paragraph(
            "※ 미국 EIA Form EIA-860M(월별 운영 배터리 저장 설문)에서 실시간 수집한 월별 운영용량 "
            "추이입니다. GWh는 평균 4시간 duration 가정의 추정치이며, 매월 신규 데이터로 갱신됩니다."
        )
        for _r in _p_eianote.runs:
            _r.font.size = Pt(8); _r.font.name = FONT
            _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        _p_eiaerr = doc.add_paragraph(
            f"ℹ️ 미국 EIA 라이브 용량 추이는 현재 표시할 수 없습니다 ({_err_eia}). "
            "EIA_API_KEY 환경변수가 설정된 배포 환경에서 자동 활성화되며, 그 전까지는 "
            "본 보고서 2장(시장별 심층 분석)의 스냅샷 용량 수치를 참고하십시오."
        )
        for _r in _p_eiaerr.runs:
            _r.font.size = Pt(9); _r.font.name = FONT

    # 9. 시나리오 분석
    _h9 = doc.add_heading("9. 시나리오 분석", level=1)
    _h9.paragraph_format.page_break_before = True
    _add_bookmark(_h9, "_sec9")
    doc.add_paragraph("각국 시장 및 거시경제 상황에 따른 BESS 확산 중장기 시나리오 전망입니다.")

    # Scenario comparison chart
    _h2b = doc.add_heading("9.1 시나리오별 용량 전망 비교", level=2)
    _add_bookmark(_h2b, "_sec9_1")
    _add_chart_to_doc(doc, _chart_scenario_comparison())

    # Scenario table
    _h2b = doc.add_heading("9.2 시나리오별 수치 비교", level=2)
    _add_bookmark(_h2b, "_sec9_2")
    scen_rows = []
    scen_rows.append(["연도", "보수적", "기준", "낙관적"])
    for yr in md.YEARS:
        cons = md.SCENARIOS["보수적 (Conservative)"]["capacity_gwh"].get(yr, 0)
        base = md.SCENARIOS["기본 (Base)"]["capacity_gwh"].get(yr, 0)
        opti = md.SCENARIOS["낙관적 (Optimistic)"]["capacity_gwh"].get(yr, 0)
        scen_rows.append([str(yr), f"{cons} GWh", f"{base} GWh", f"{opti} GWh"])
    _styled_table(doc, scen_rows[0], scen_rows[1:], col_widths_mm=[40, 40, 40, 40])

    # Section 9 interpretation
    _sb = md.SCENARIOS["기본 (Base)"]
    _sc = md.SCENARIOS["보수적 (Conservative)"]
    _so = md.SCENARIOS["낙관적 (Optimistic)"]
    _scen_years = sorted(_sb["capacity_gwh"].keys())
    _scen_start, _scen_end = _scen_years[0], _scen_years[-1]
    _p_scen = doc.add_paragraph(
        f"기본 시나리오({_sb['description']}) 기준 {_scen_start}~{_scen_end}년 CAGR {_sb['cagr_pct']}%로 "
        f"{_scen_end}년 {_sb['capacity_gwh'].get(_scen_end, 0)} GWh"
        f"(시장 가치 ${_sb['market_value_b'].get(_scen_end, 0)}B)가 전망됩니다. "
        f"낙관 시나리오({_so['description']}) 실현 시 CAGR {_so['cagr_pct']}%로 "
        f"{_scen_end}년 {_so['capacity_gwh'].get(_scen_end, 0)} GWh까지 확대 가능하며, "
        f"보수 시나리오({_sc['description']}) 하에서도 "
        f"{_sc['capacity_gwh'].get(_scen_end, 0)} GWh(CAGR {_sc['cagr_pct']}%)로 "
        "견조한 성장이 예상됩니다. "
        "시나리오 편차의 핵심 변수는 IRA·RE정책 지속성, 배터리 가격 하락 속도, "
        "전력시장 설계 개혁의 실행 속도입니다."
    )
    _p_scen.paragraph_format.space_before = Pt(4)
    for _r in _p_scen.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # ============================================================
    # 10. BESS 사업 개발 및 투자 분석
    # ============================================================
    _h10 = doc.add_heading("10. BESS 사업 개발 및 투자 분석", level=1)
    _h10.paragraph_format.page_break_before = True
    _add_bookmark(_h10, "_sec10")
    doc.add_paragraph(
        "BESS 프로젝트의 수익 모델(Revenue Stacking), 투자 경제성(IRR/Payback), "
        "Offtake/PPA 계약 구조를 시장별로 분석합니다. 사업 개발 단계에서의 핵심 의사결정 요소를 제시합니다."
    )

    # 10.1 시장별 수익 스태킹 분석
    _h2b = doc.add_heading("10.1 시장별 수익 스태킹(Revenue Stacking) 분석", level=2)
    _add_bookmark(_h2b, "_sec10_1")
    for mkt_name in ["미국", "영국", "호주", "한국", "EU"]:
        mkt_rs = md.REVENUE_STACKING.get(mkt_name, {})
        if not mkt_rs:
            continue
        doc.add_heading(f"  ■ {mkt_name}", level=3)
        rs_headers = ["수익원", "비중(%)", "평균수익($/kWh/yr)", "추세", "설명"]
        rs_rows = []
        total_rev = 0
        for _key, _val in mkt_rs.items():
            rs_rows.append([
                _key.replace("_", " ").title(),
                str(_val["share_pct"]),
                str(_val["avg_revenue_kwh_yr"]),
                _val["trend"],
                _val["desc"][:120] + "…" if len(_val["desc"]) > 120 else _val["desc"],
            ])
            total_rev += _val["avg_revenue_kwh_yr"] * _val["share_pct"] / 100
        _styled_table(doc, rs_headers, rs_rows, col_widths_mm=[24, 12, 18, 10, 96])
        _p_rs = doc.add_paragraph(
            f"→ {mkt_name} 시장 가중평균 예상 수익: ${total_rev:.0f}/kWh/yr (Revenue Stacking 기준)"
        )
        _p_rs.paragraph_format.space_before = Pt(2)
        for _r in _p_rs.runs:
            _r.font.size = Pt(9); _r.font.name = FONT; _r.bold = True

    # 10.2 투자 경제성 비교
    _h2b = doc.add_heading("10.2 프로젝트 유형별 투자 경제성 비교", level=2)
    _add_bookmark(_h2b, "_sec10_2")
    inv_headers = ["프로젝트 유형", "CAPEX\n($/kWh)", "OPEX\n($/kWh/yr)", "IRR\n(Base)", "IRR\n(Opt.)",
                   "Payback\n(년)", "수명\n(년)", "LCOS\n($/kWh)"]
    inv_rows = []
    for _k, _v in md.INVESTMENT_ECONOMICS.items():
        inv_rows.append([
            _v["name"],
            str(_v["capex_per_kwh"]),
            str(_v["opex_per_kwh_yr"]),
            f"{_v['irr_base_pct']}%",
            f"{_v['irr_optimistic_pct']}%",
            str(_v["payback_years"]),
            str(_v["project_life_years"]),
            f"${_v['lcoe_kwh']:.2f}",
        ])
    _styled_table(doc, inv_headers, inv_rows, col_widths_mm=[36, 16, 16, 14, 14, 14, 12, 38])
    _p_inv = doc.add_paragraph(
        "유틸리티급 4시간 BESS는 Base Case IRR 12.5%로 안정적이며, 2시간 피크 대응형은 "
        "짧은 duration에도 불구하고 높은 보조서비스 수익으로 14.0% IRR 달성이 가능합니다. "
        "태양광+ESS 하이브리드는 LCOS $0.10/kWh로 가장 낮은 균등화 저장 비용을 보여 "
        "IRA ITC 적용 시 경쟁력이 극대화됩니다."
    )
    for _r in _p_inv.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 10.3 Offtake / PPA 구조 비교
    _h2b = doc.add_heading("10.3 Offtake 및 PPA 계약 구조 비교", level=2)
    _add_bookmark(_h2b, "_sec10_3")
    oft_headers = ["계약 유형", "계약기간", "리스크", "수익 확실성", "주요 시장", "설명"]
    oft_rows = []
    for _o in md.OFFTAKE_STRUCTURES:
        oft_rows.append([
            _o["type"], _o["duration_yr"], _o["risk_profile"],
            _o["revenue_certainty"], _o["typical_market"],
            _o["desc"][:100] + "…" if len(_o["desc"]) > 100 else _o["desc"],
        ])
    _styled_table(doc, oft_headers, oft_rows, col_widths_mm=[24, 14, 12, 14, 18, 78])
    _p_oft = doc.add_paragraph(
        "프로젝트 파이낸싱 관점에서 Tolling Agreement와 PPA가 가장 유리하며, "
        "비소구금융(Non-recourse Project Finance) 조달 시 장기 계약(10년+)이 필수적입니다. "
        "영국 시장은 Merchant 모델이 주류이나 최근 Contracted+Merchant 혼합 구조로 전환하는 추세입니다. "
        "한국은 전력시장 구조상 한전/KPX 기반 계약이 중심이나, 전력시장 개편 시 다양한 계약 구조 도입이 예상됩니다."
    )
    for _r in _p_oft.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # ============================================================
    # 11. 전력시장 및 거래 동향
    # ============================================================
    _h11 = doc.add_heading("11. 전력시장 및 거래 동향", level=1)
    _h11.paragraph_format.page_break_before = True
    _add_bookmark(_h11, "_sec11")
    doc.add_paragraph(
        "주요 시장별 전력 거래 구조, BESS 참여 메커니즘, 수익 기회를 분석합니다. "
        "각국 전력시장 설계(Market Design)의 차이가 BESS 사업성에 미치는 영향을 심층 조망합니다."
    )

    for pm_name, pm_data in md.POWER_MARKET_STRUCTURES.items():
        doc.add_heading(f"11.{list(md.POWER_MARKET_STRUCTURES.keys()).index(pm_name)+1} {pm_name} 전력시장", level=2)
        pm_rows = [
            ["시장 유형", pm_data["market_type"]],
            ["주요 시장/플랫폼", ", ".join(pm_data["key_markets"])],
            ["결제 구조", pm_data["settlement"]],
            ["BESS 참여 범위", pm_data["bess_participation"]],
            ["평균 스프레드", pm_data["avg_spread_kwh"]],
        ]
        _styled_table(doc, ["항목", "내용"], pm_rows, col_widths_mm=[35, 125])
        _p_pm = doc.add_paragraph(f"[동향] {pm_data['key_trend']}")
        _p_pm.paragraph_format.space_before = Pt(4)
        for _r in _p_pm.runs:
            _r.font.size = Pt(12); _r.font.name = FONT; _r.italic = True

    # 전력 거래 전략 시사점
    _h2b = doc.add_heading("11.6 BESS 전력 거래 전략 시사점", level=2)
    _add_bookmark(_h2b, "_sec11_6")
    _trading_insights = [
        "Revenue Stacking 필수화: 단일 수익원 의존은 수익성 리스크가 높아 복수 시장 동시 참여(에너지+보조서비스+용량) 전략이 표준화되고 있습니다.",
        "AI/ML 기반 입찰 최적화: Tesla Autobidder, Fluence Mosaic 등 AI 기반 EMS가 수동 입찰 대비 15-25% 수익 향상을 달성하고 있어, "
        "소프트웨어 역량이 BESS 사업의 핵심 차별화 요소로 부상하고 있습니다.",
        "Duration 장기화 트렌드: 2시간→4시간→8시간으로 저장 시간이 확대되는 추세. 장기 저장(LDES)은 용량시장 및 계통 안정성 측면에서 프리미엄 수익 가능.",
        "시장 간 수익 다변화: 단일 국가/시장 의존 리스크 분산을 위해 복수 시장 포트폴리오 전략이 중요합니다. "
        "미국(ITC 수혜)+영국(보조서비스)+호주(FCAS) 조합이 최적 분산 포트폴리오로 평가됩니다.",
        "한국 전력시장 개편 대비: Cost-Based → Bid-Based Pool 전환 시 BESS 에너지 차익거래 기회가 비약적으로 확대될 전망. "
        "실시간 시장 도입에 대비한 입찰 전략 및 EMS 역량 확보가 시급합니다.",
    ]
    for _ti in _trading_insights:
        _p_ti = doc.add_paragraph(_ti, style="List Bullet")
        for _r in _p_ti.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    # ============================================================
    # 12. BESS 운영 및 자산관리
    # ============================================================
    _h12 = doc.add_heading("12. BESS 운영 및 자산관리", level=1)
    _h12.paragraph_format.page_break_before = True
    _add_bookmark(_h12, "_sec12")
    doc.add_paragraph(
        "BESS 프로젝트의 장기 운영 성능, O&M 비용 추이, 배터리 열화 관리, "
        "에너지관리시스템(EMS) 플랫폼을 분석합니다. 20년 이상 프로젝트 수명 동안의 "
        "자산 가치 극대화 전략을 제시합니다."
    )

    # 12.1 핵심 성능 지표
    _h2b = doc.add_heading("12.1 핵심 성능 지표(KPI)", level=2)
    _add_bookmark(_h2b, "_sec12_1")
    perf = md.OPERATIONS_DATA["performance_metrics"]
    perf_rows = []
    for pk, pv in perf.items():
        perf_rows.append([pk.replace("_", " ").title(), pv["value"], pv["trend"], pv["desc"]])
    _styled_table(doc, ["지표", "기준값", "추세", "설명"], perf_rows, col_widths_mm=[26, 16, 10, 108])

    # 12.2 O&M 비용 추이
    _h2b = doc.add_heading("12.2 O&M 비용 추이", level=2)
    _add_bookmark(_h2b, "_sec12_2")
    om = md.OPERATIONS_DATA["om_cost_trends"]
    om_headers = ["연도", "고정비($/kW/yr)", "변동비($/MWh)", "총 O&M($/kWh/yr)"]
    om_rows = []
    for oyr in sorted(om.keys()):
        od = om[oyr]
        om_rows.append([str(oyr), str(od["fixed_per_kw_yr"]), str(od["variable_per_mwh"]), str(od["total_per_kwh_yr"])])
    _styled_table(doc, om_headers, om_rows, col_widths_mm=[22, 46, 46, 46])
    _first_om = list(sorted(om.keys()))[0]
    _last_om = list(sorted(om.keys()))[-1]
    _om_drop = round((1 - om[_last_om]["total_per_kwh_yr"] / om[_first_om]["total_per_kwh_yr"]) * 100)
    _p_om = doc.add_paragraph(
        f"O&M 비용은 {_first_om}년 ${om[_first_om]['total_per_kwh_yr']}/kWh/yr에서 "
        f"{_last_om}년 ${om[_last_om]['total_per_kwh_yr']}/kWh/yr로 {_om_drop}% 감소하였습니다. "
        "원격 모니터링 확대, 예측정비(Predictive Maintenance) 도입, 규모의 경제가 주요 원인입니다."
    )
    for _r in _p_om.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # 12.3 EMS 플랫폼 비교
    _h2b = doc.add_heading("12.3 에너지관리시스템(EMS) 플랫폼 비교", level=2)
    _add_bookmark(_h2b, "_sec12_3")
    ems = md.OPERATIONS_DATA["ems_platforms"]
    ems_headers = ["플랫폼", "벤더", "핵심 기능", "주요 시장"]
    ems_rows = [[e["name"], e["vendor"], e["feature"], e["market"]] for e in ems]
    _styled_table(doc, ems_headers, ems_rows, col_widths_mm=[24, 22, 78, 36])

    # 12.4 배터리 열화 관리 전략
    _h2b = doc.add_heading("12.4 배터리 열화 관리 및 수명 연장 전략", level=2)
    _add_bookmark(_h2b, "_sec12_4")
    for _dm in md.OPERATIONS_DATA["degradation_mgmt"]:
        _p_dm = doc.add_paragraph(_dm, style="List Bullet")
        for _r in _p_dm.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    _p_dm_sum = doc.add_paragraph(
        "체계적인 열화 관리는 BESS 프로젝트의 20년 수명 동안 누적 수익을 15-25% 향상시킬 수 있습니다. "
        "특히 Augmentation 시점과 규모의 최적화가 프로젝트 NPV에 가장 큰 영향을 미치며, "
        "초기 설계 단계에서 증설 공간(Bay) 확보가 필수적입니다."
    )
    for _r in _p_dm_sum.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # ============================================================
    # 13. 안전·화재 및 규제 기준
    # ============================================================
    _h13 = doc.add_heading("13. 안전·화재 및 규제 기준", level=1)
    _h13.paragraph_format.page_break_before = True
    _add_bookmark(_h13, "_sec13")
    doc.add_paragraph(
        "BESS 안전은 프로젝트 인허가, 보험, 금융 조달의 핵심 요건입니다. "
        "주요국 안전 규격, 화재 사고 사례 및 교훈, 설계 반영 사항을 분석합니다."
    )

    _h2b = doc.add_heading("13.1 글로벌 ESS 안전 규격 비교", level=2)
    _add_bookmark(_h2b, "_sec13_1")
    sf_headers = ["규격", "적용 지역", "범위", "핵심 요구사항"]
    sf_rows = [[s["standard"], s["region"], s["scope"], s["key_req"]] for s in md.SAFETY_STANDARDS]
    _styled_table(doc, sf_headers, sf_rows, col_widths_mm=[20, 18, 26, 96])

    _h2b = doc.add_heading("13.2 주요 ESS 화재 사고 사례 및 교훈", level=2)
    _add_bookmark(_h2b, "_sec13_2")
    fi_headers = ["연도", "위치", "원인", "피해", "교훈"]
    fi_rows = [[str(f["year"]), f["location"], f["cause"], f["damage"], f["lesson"]] for f in md.FIRE_INCIDENTS]
    _styled_table(doc, fi_headers, fi_rows, col_widths_mm=[10, 20, 24, 24, 82])
    _p_fi = doc.add_paragraph(
        "ESS 화재 사고는 산업 전체의 안전 규제 강화를 촉발하고 있습니다. "
        "열폭주 전파 방지 설계(셀 간/모듈 간 방화벽), 가스 감지 및 자동 소화 시스템, "
        "BMS 이중화, 정기 점검 프로토콜이 필수 설계 요소로 자리잡고 있습니다. "
        "보험사들도 UL 9540A 시험 결과 및 소방 설계 도면을 대출/보험 심사 필수 서류로 요구하는 추세입니다."
    )
    for _r in _p_fi.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("13.3 안전 설계 체크리스트", level=2)
    _add_bookmark(_h2b, "_sec13_3")
    _safety_checklist = [
        "셀 레벨: UL 9540A 열폭주 시험 통과, IEC 62619 인증, 셀 간 방열 패드/에어갭",
        "모듈/랙 레벨: 모듈 간 방화벽, 열폭주 전파 방지 설계, 가스 벤트 경로 확보",
        "컨테이너 레벨: 가스 감지(CO, H₂, VOC), 자동 소화 시스템(에어로졸/워터미스트), 폭발 방지 환기팬",
        "시스템 레벨: BMS 이중화(Main+Backup), PCS 절연 모니터링, 접지 결함 감지(GFD)",
        "사이트 레벨: 이격거리(NFPA 855 기준), 소방차 접근로, 비상 차단 스위치(E-Stop), 소방서 사전 훈련",
        "운영 레벨: 24/7 원격 모니터링, 이상 징후 AI 분석, 분기별 안전 점검, 연 1회 종합 안전 감사",
    ]
    for _sc in _safety_checklist:
        _p_sc = doc.add_paragraph(_sc, style="List Bullet")
        for _r in _p_sc.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    # ============================================================
    # 14. 배터리 기술 동향 및 차세대 기술
    # ============================================================
    _h14 = doc.add_heading("14. 배터리 기술 동향 및 차세대 기술", level=1)
    _h14.paragraph_format.page_break_before = True
    _add_bookmark(_h14, "_sec14")
    doc.add_paragraph(
        "현재 주류 기술(LFP/NMC)과 차세대 기술(Na-ion, 전고체, VRFB, 철-공기)의 "
        "성능·비용·상용화 전망을 비교 분석합니다. 장기 에너지 저장(LDES) 시장 동향도 포함합니다."
    )

    _h2b = doc.add_heading("14.1 배터리 기술 비교표", level=2)
    _add_bookmark(_h2b, "_sec14_1")
    _battery_techs = md.get_battery_technologies()
    bt_headers = ["기술", "에너지밀도\n(Wh/kg)", "사이클수명", "비용\n($/kWh)", "상용화 단계"]
    bt_rows = [[b["tech"], b["energy_density_wh_kg"], b["cycle_life"], b["cost_per_kwh"], b["status"]] for b in _battery_techs]
    _styled_table(doc, bt_headers, bt_rows, col_widths_mm=[24, 20, 20, 18, 78])

    _h2b = doc.add_heading("14.2 기술별 상세 분석", level=2)
    _add_bookmark(_h2b, "_sec14_2")
    for bt in _battery_techs:
        doc.add_heading(f"  ■ {bt['tech']}", level=3)
        bt_detail = [
            ["화학식", bt["chemistry"]],
            ["장점", bt["pros"]],
            ["단점", bt["cons"]],
            ["전망", bt["outlook"]],
        ]
        _styled_table(doc, ["항목", "내용"], bt_detail, col_widths_mm=[25, 135])

    _h2b = doc.add_heading("14.3 장기 에너지 저장(LDES) 시장 전망", level=2)
    _add_bookmark(_h2b, "_sec14_3")
    _ldes = md.LDES_MARKET
    _p_ldes = doc.add_paragraph(
        f"LDES 시장은 {_ldes['base_size_gwh']} GWh({_ldes['base_year']}) → "
        f"{_ldes['target_size_gwh']} GWh({_ldes['target_year']})로 "
        f"CAGR {_ldes['cagr_pct']}% 성장 전망. "
        f"정의: {_ldes['definition']}."
    )
    for _r in _p_ldes.runs:
        _r.font.size = Pt(12); _r.font.name = FONT; _r.bold = True
    for _d in _ldes["key_drivers"]:
        _p_ld = doc.add_paragraph(_d, style="List Bullet")
        for _r in _p_ld.runs:
            _r.font.size = Pt(12); _r.font.name = FONT
    ldes_headers = ["기술", "시장 점유율(%)", "핵심 장점"]
    ldes_rows = [[t["tech"], str(t["share_pct"]), t["advantage"]] for t in _ldes["competing_techs"]]
    _styled_table(doc, ldes_headers, ldes_rows, col_widths_mm=[35, 25, 100])

    # ============================================================
    # 15. 인허가 및 사업 개발 프로세스
    # ============================================================
    _h15 = doc.add_heading("15. 인허가 및 사업 개발 프로세스", level=1)
    _h15.paragraph_format.page_break_before = True
    _add_bookmark(_h15, "_sec15")
    doc.add_paragraph(
        "시장별 인허가 절차, 그리드 연결 대기시간, 토지 요건 등 "
        "BESS 프로젝트 개발 실무에 필수적인 정보를 정리합니다."
    )

    for pm_idx, (pm_name, pm_data) in enumerate(md.PERMITTING_DATA.items()):
        doc.add_heading(f"15.{pm_idx+1} {pm_name}", level=2)
        pm_rows = [
            ["총 인허가 기간", f"{pm_data['total_timeline_months']}개월"],
            ["그리드 연결 대기", f"{pm_data['grid_connection_wait_months']}개월"],
            ["필요 토지", f"{pm_data['land_req_acre_per_mwh']} acre/MWh"],
            ["필요 인허가", ", ".join(pm_data["key_permits"])],
        ]
        _styled_table(doc, ["항목", "내용"], pm_rows, col_widths_mm=[35, 125])
        _p_gc = doc.add_paragraph(f"[그리드 연결 현황] {pm_data['grid_challenge']}")
        _p_gc.paragraph_format.space_before = Pt(4)
        for _r in _p_gc.runs:
            _r.font.size = Pt(12); _r.font.name = FONT; _r.italic = True
        _p_tip = doc.add_paragraph(f"[실무 Tip] {pm_data['tips']}")
        for _r in _p_tip.runs:
            _r.font.size = Pt(12); _r.font.name = FONT; _r.bold = True

    # ============================================================
    # 16. 프로젝트 파이낸싱
    # ============================================================
    _h16 = doc.add_heading("16. 프로젝트 파이낸싱", level=1)
    _h16.paragraph_format.page_break_before = True
    _add_bookmark(_h16, "_sec16")
    doc.add_paragraph(
        "BESS 프로젝트의 자금 조달 구조, 금융기관 심사 요건(Bankability), "
        "보험 요건을 분석합니다. 프로젝트 파이낸스(PF) 시장의 최신 동향을 반영합니다."
    )

    _h2b = doc.add_heading("16.1 자금 조달 구조 비교", level=2)
    _add_bookmark(_h2b, "_sec16_1")
    pf_headers = ["유형", "레버리지", "기간(년)", "최소 DSCR", "주요 대출기관"]
    pf_rows = []
    for pfs in md.PROJECT_FINANCING["structures"]:
        pf_rows.append([pfs["type"], pfs["leverage"], pfs["tenor_yr"],
                        pfs["min_dscr"], pfs["key_lenders"][:70] + "…" if len(pfs["key_lenders"]) > 70 else pfs["key_lenders"]])
    _styled_table(doc, pf_headers, pf_rows, col_widths_mm=[28, 14, 12, 16, 90])

    for pfs in md.PROJECT_FINANCING["structures"]:
        _p_pfs = doc.add_paragraph(f"■ {pfs['type']}: {pfs['desc']}")
        _p_pfs.paragraph_format.space_before = Pt(4)
        for _r in _p_pfs.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("16.2 Bankability 요건 (금융기관 심사 핵심)", level=2)
    _add_bookmark(_h2b, "_sec16_2")
    for _bk in md.PROJECT_FINANCING["bankability_requirements"]:
        _p_bk = doc.add_paragraph(_bk, style="List Bullet")
        for _r in _p_bk.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("16.3 보험 요건", level=2)
    _add_bookmark(_h2b, "_sec16_3")
    ins_headers = ["보험 유형", "보장 내용", "요율/한도"]
    ins_rows = [[i["type"], i["desc"], i["rate"]] for i in md.PROJECT_FINANCING["insurance_coverage"]]
    _styled_table(doc, ins_headers, ins_rows, col_widths_mm=[28, 88, 44])

    # ============================================================
    # 17. EPC 계약 구조
    # ============================================================
    _h17 = doc.add_heading("17. EPC 계약 구조", level=1)
    _h17.paragraph_format.page_break_before = True
    _add_bookmark(_h17, "_sec17")
    doc.add_paragraph(
        "BESS 프로젝트의 EPC 계약 유형(Turnkey/EPCM/Split/BOO), "
        "핵심 상업 조건(Performance Guarantee, LD, Warranty), 비용 구조를 분석합니다."
    )

    _h2b = doc.add_heading("17.1 EPC 계약 유형 비교", level=2)
    _add_bookmark(_h2b, "_sec17_1")
    epc_headers = ["계약 유형", "리스크 부담", "가격 구조", "장점", "단점"]
    epc_rows = []
    for ec in md.EPC_CONTRACT_DATA["contract_types"]:
        epc_rows.append([ec["type"], ec["risk_owner"], ec["price_structure"],
                         ec["pros"][:70] + "…" if len(ec["pros"]) > 70 else ec["pros"],
                         ec["cons"][:70] + "…" if len(ec["cons"]) > 70 else ec["cons"]])
    _styled_table(doc, epc_headers, epc_rows, col_widths_mm=[24, 18, 20, 49, 49])

    for ec in md.EPC_CONTRACT_DATA["contract_types"]:
        _p_ec = doc.add_paragraph(f"■ {ec['type']}: {ec['desc']}")
        _p_ec.paragraph_format.space_before = Pt(4)
        for _r in _p_ec.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("17.2 핵심 상업 조건", level=2)
    _add_bookmark(_h2b, "_sec17_2")
    ct_headers = ["조건", "상세 내용"]
    ct_rows = [[ct["term"], ct["desc"]] for ct in md.EPC_CONTRACT_DATA["key_commercial_terms"]]
    _styled_table(doc, ct_headers, ct_rows, col_widths_mm=[35, 125])

    _h2b = doc.add_heading("17.3 BESS 프로젝트 비용 구조(Cost Breakdown)", level=2)
    _add_bookmark(_h2b, "_sec17_3")
    cb = md.EPC_CONTRACT_DATA["cost_breakdown"]
    cb_headers = ["항목", "비중(%)", "포함 내용"]
    cb_rows = sorted(
        [[k.replace("_", " ").title(), str(v["share_pct"]), v["desc"]] for k, v in cb.items()],
        key=lambda x: int(x[1]), reverse=True
    )
    _styled_table(doc, cb_headers, cb_rows, col_widths_mm=[35, 18, 107])
    _p_cb = doc.add_paragraph(
        "배터리 셀이 전체 CAPEX의 40%를 차지하여 셀 조달 전략이 프로젝트 경제성의 핵심입니다. "
        "EPC 마진 10%는 Full Turnkey 기준이며, EPCM/Split 방식 시 5-8%로 절감 가능합니다."
    )
    for _r in _p_cb.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    # ============================================================
    # 18. 전문가 종합 의견 및 전략적 시사점 (기존 13장 → 18장)
    # ============================================================
    _h18 = doc.add_heading("18. 전문가 종합 의견 및 전략적 시사점", level=1)
    _h18.paragraph_format.page_break_before = True
    _add_bookmark(_h18, "_sec18")

    _h2b = doc.add_heading("18.1 시장 전망 종합", level=2)
    _add_bookmark(_h2b, "_sec18_1")
    _p_exp1 = doc.add_paragraph(
        f"글로벌 BESS 시장은 {_yr}년 기준 {md.GLOBAL_CAPACITY_GWH.get(_yr, 0)} GWh 규모로 "
        f"{_yr - 1}년 대비 {_yoy_g}% 성장하였으며, {md.YEARS[-1]}년 "
        f"{md.GLOBAL_CAPACITY_GWH[md.YEARS[-1]]} GWh까지 확대될 전망입니다. "
        "이는 재생에너지 확대에 따른 계통 유연성 확보 수요, 배터리 셀 가격의 지속적 하락, "
        "주요국의 에너지 전환 정책이 복합적으로 작용한 구조적 성장입니다. "
        "특히 미국 IRA, EU REPowerEU, 영국 Capacity Market 등 정책 프레임워크가 "
        "BESS 투자의 확실성을 높이고 있습니다."
    )
    for _r in _p_exp1.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("18.2 가격 전망 및 경제성 분석", level=2)
    _add_bookmark(_h2b, "_sec18_2")
    # LFP $30 이하 진입 연도 동적 추출
    _lfp_30_yr = next((y for y in md.YEARS if md.LFP_CELL_PRICE.get(y, 999) <= 30), None)
    _lfp_final_price = md.LFP_CELL_PRICE[_last_yr]
    if _lfp_30_yr:
        _price_outlook = (
            f"{_lfp_30_yr}년경 $30/kWh 수준에 도달할 것으로 전망되며, "
            f"{_last_yr}년에는 ${_lfp_final_price}/kWh 수준이 예상됩니다."
        )
    else:
        _price_outlook = (
            f"{_last_yr}년 ${_lfp_final_price}/kWh 수준이 전망되며, "
            "$30/kWh 도달은 데이터 범위 외에서 예상됩니다."
        )
    _p_exp2 = doc.add_paragraph(
        f"LFP 셀 가격은 {md.LFP_CELL_PRICE[_first_yr]}→{md.LFP_CELL_PRICE[_last_yr]} $/kWh "
        f"({abs(_lfp_cagr)}% CAGR 하락)로 빠르게 하락하고 있으며, "
        f"{_price_outlook} "
        "이 가격대에서 BESS는 가스 피커(Gas Peaker) 대비 완전한 경제적 우위를 확보하게 되며, "
        "보조금 없이도 자립적 사업성이 확보되는 '그리드 패리티' 달성이 예상됩니다. "
        "다만 리튬 가격 바닥 도달, 공급망 재편 비용, 인플레이션 등 하방 리스크에도 주의해야 합니다."
    )
    for _r in _p_exp2.runs:
        _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("18.3 EPC 사업 전략 시사점", level=2)
    _add_bookmark(_h2b, "_sec18_3")
    _epc_insights = [
        "셀 소싱 전략: LFP 중심의 멀티 공급사 전략으로 리스크 분산 및 원가 경쟁력 확보가 필수입니다. "
        "CATL/BYD/EVE 등 중국 Tier 1과 Samsung SDI/LG 등 한국 제조사의 이중 소싱을 권장합니다.",
        "시장 진입 우선순위: 미국(IRA ITC 30%+), 호주(FCAS 수익), 영국(Capacity Market)을 "
        "1차 타겟 시장으로 권장하며, 각 시장의 인허가 리드타임과 그리드 연결 대기 시간을 "
        "사전에 분석하여 프로젝트 일정에 반영해야 합니다.",
        "기술 차별화: 4시간 이상 장기저장(LDES), AI 기반 에너지 관리 시스템(EMS), "
        "모듈형 설계를 통한 확장성 확보가 경쟁 우위의 핵심입니다.",
        "리스크 관리: 환율 헤지, 원자재 가격 연동 계약(Pass-through), "
        "성능 보증(Performance Guarantee) 구조 설계를 통해 프로젝트 수익성을 보호해야 합니다.",
        "안전 규정 대응: NFPA 855, IEC 62933, UL 9540A 등 글로벌 안전 인증을 "
        "사전 확보하고, 화재 시뮬레이션 및 소방 설계를 프로젝트 초기 단계에 반영해야 합니다.",
    ]
    for insight in _epc_insights:
        _p_ins = doc.add_paragraph(insight, style="List Bullet")
        for _r in _p_ins.runs:
            _r.font.size = Pt(12); _r.font.name = FONT

    _h2b = doc.add_heading("18.4 핵심 리스크 요인", level=2)
    _add_bookmark(_h2b, "_sec18_4")
    _risks = [
        ("지정학적 리스크", "미·중 기술 갈등, 수출 통제, 관세 부과 → 셀 공급 다변화 필수"),
        ("정책 불확실성", "IRA 지속 여부, EU 보조금 변경 → 시나리오별 수익 모델링 강화"),
        ("그리드 연결 지연", "영국·호주 등 인프라 병목 → 조기 그리드 연결 신청 및 대체 부지 확보"),
        ("안전·화재 리스크", "ESS 화재 사고 지속 → 최신 안전 기준 준수, 열폭주 방지 설계 강화"),
        ("기술 진부화", "Na-ion, 고체 전해질 등 차세대 기술 등장 → 기술 로드맵 지속 모니터링"),
    ]
    risk_rows = [[r[0], r[1]] for r in _risks]
    _styled_table(doc, ["리스크 요인", "대응 전략"], risk_rows, col_widths_mm=[40, 120])

    # Disclaimer
    doc.add_paragraph()
    _p_disc = doc.add_paragraph(
        "※ 본 보고서는 공개 데이터 및 RSS 뉴스 피드를 기반으로 AI 에이전트가 자동 생성한 참고 자료입니다. "
        "투자 또는 사업 의사결정 시 별도의 전문가 검증을 권장합니다."
    )
    _p_disc.paragraph_format.space_before = Pt(12)
    for _r in _p_disc.runs:
        _r.font.size = Pt(9); _r.font.name = FONT
        _r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(out_path)
    return out_path

def generate_pdf_report():
    """Generate PDF — Linux: LibreOffice headless, Windows: docx2pdf(COM)."""
    import shutil
    import subprocess

    word_path = generate_word_report()
    pdf_path = word_path.replace('.docx', '.pdf')
    tmp_dir = tempfile.gettempdir()

    if platform.system() != "Windows":
        # Linux / HF Spaces: LibreOffice headless 변환
        tmp_docx = os.path.join(tmp_dir, "bess_report_tmp.docx")
        shutil.copy2(word_path, tmp_docx)
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", tmp_dir, tmp_docx],
                capture_output=True, text=True, timeout=180,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice 변환 실패: {result.stderr[:300]}")
            tmp_pdf = os.path.join(tmp_dir, "bess_report_tmp.pdf")
            if os.path.exists(tmp_pdf):
                shutil.copy2(tmp_pdf, pdf_path)
                for p in (tmp_docx, tmp_pdf):
                    try:
                        os.unlink(p)
                    except Exception:
                        pass
                return os.path.abspath(pdf_path)
            raise FileNotFoundError("LibreOffice PDF 파일이 생성되지 않았습니다.")
        except FileNotFoundError:
            raise RuntimeError(
                "LibreOffice가 설치되어 있지 않습니다. "
                "packages.txt에 'libreoffice-writer'를 추가하고 Space를 재빌드하세요."
            )
    else:
        # Windows: docx2pdf (COM/Microsoft Word)
        import threading
        import docx2pdf
        tmp_docx = os.path.join(tmp_dir, "bess_report_tmp.docx")
        tmp_pdf = os.path.join(tmp_dir, "bess_report_tmp.pdf")
        shutil.copy2(word_path, tmp_docx)

        errors = []
        def _convert():
            try:
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    docx2pdf.convert(tmp_docx, tmp_pdf)
                finally:
                    pythoncom.CoUninitialize()
            except Exception as e:
                errors.append(e)

        t = threading.Thread(target=_convert, daemon=True)
        t.start()
        t.join(timeout=120)

        if errors:
            raise RuntimeError(f"PDF 변환 오류: {errors[0]}")

        if os.path.exists(tmp_pdf):
            shutil.copy2(tmp_pdf, pdf_path)
            for p in (tmp_docx, tmp_pdf):
                try:
                    os.unlink(p)
                except Exception:
                    pass
            return os.path.abspath(pdf_path)
        raise FileNotFoundError(f"PDF가 생성되지 않았습니다: {pdf_path}")
